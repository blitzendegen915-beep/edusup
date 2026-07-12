"""API不使用部分の回帰テスト。実行: python -m exam_app.tests.test_offline"""
import json
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from exam_app import checks, build_docx, extract  # noqa: E402


def _draft(**exam):
    return {
        "exam": {"title": "テスト試験", **exam},
        "sections": [
            {"no": 1, "type": "fill_blank", "points_each": 1, "count": 2, "questions": [
                {"number": 1, "body": "Many people (  ) to Nagoya.",
                 "answer": "commute", "source_ref": "L3 Part3①(6)", "alt_answer_risk": "低"},
                {"number": 2, "body": "Hope for peace is ( u ).",
                 "answer": "universal", "source_ref": "L4 Part4①(6)", "alt_answer_risk": "uヒント"},
            ]},
        ],
    }


def test_checks_pass():
    assert checks.run_all(_draft(written_points=2)) == []


def test_checks_detect():
    d = _draft(written_points=5)                      # 配点ズレ
    d["sections"][0]["questions"][1]["number"] = 3    # 連番ズレ
    d["sections"][0]["questions"][1]["source_ref"] = ""  # 出典なし
    issues = checks.run_all(d)
    assert any("配点" in i for i in issues)
    assert any("連番" in i for i in issues)
    assert any("出典なし" in i for i in issues)


def test_checks_duplicates_and_bias():
    d = _draft()
    d["sections"].append({"no": 2, "type": "choice", "points_each": 2, "count": 3,
        "questions": [
            {"number": i, "body": f"C{i}", "answer": "ア",
             "source_ref": f"s{i}", "alt_answer_risk": ""} for i in (1, 2, 3)]})
    issues = checks.run_all(d)
    assert any("偏って" in i for i in issues)
    assert not any("'ア' が" in i for i in issues)  # 選択記号は重複扱いしない
    d["sections"][0]["questions"][1]["answer"] = "commute"
    assert any("重複" in i for i in checks.run_all(d))


def test_docx_roundtrip():
    from docx import Document
    d = _draft()
    with tempfile.TemporaryDirectory() as tmp:
        files = build_docx.build_all(d, Path(tmp))
        assert [f.name for f in files] == [
            "exam_draft.docx", "answersheet_draft.docx", "modelanswer_draft.docx"]
        model = Document(str(files[2]))
        cells = [c.text for t in model.tables for r in t.rows for c in r.cells]
        assert "universal" in cells
        blank = Document(str(files[1]))
        cells = [c.text for t in blank.tables for r in t.rows for c in r.cells]
        assert "universal" not in cells  # 解答用紙に答えが漏れていない


def test_extract_three_layers():
    docs = extract.scan_folder(Path(__file__).resolve().parents[2] / "materials")
    assert docs, "materials/ が空"
    text = docs["exam_comm2_master"]["text"]
    assert "TABLE" in text or len(text) > 1000


def main():
    fns = [v for k, v in sorted(globals().items()) if k.startswith("test_")]
    for fn in fns:
        fn()
        print(f"PASS {fn.__name__}")
    print(f"\n{len(fns)}件すべて合格")


if __name__ == "__main__":
    main()
