"""docx/PDFから教材テキストを抽出する（API不使用）。

skills/exam-docx の鉄則: 段落・表・XMLの3層すべてを読む。
"""
import json
import re
import zipfile
from pathlib import Path

from docx import Document


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for ti, t in enumerate(doc.tables):
        parts.append(f"[TABLE {ti}]")
        for row in t.rows:
            cells, seen = [], set()
            for c in row.cells:
                if id(c._tc) in seen:      # 結合セルの重複を除去
                    continue
                seen.add(id(c._tc))
                cells.append(c.text.strip())
            parts.append(" | ".join(cells))
    # テキストボックス等の取りこぼしをXMLで確認
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    txbx = re.findall(r"<w:txbxContent>.*?</w:txbxContent>", xml, re.S)
    for i, box in enumerate(txbx):
        text = re.sub(r"<[^>]+>", "", box).strip()
        if text:
            parts.append(f"[TEXTBOX {i}] {text}")
    return "\n".join(parts)


def extract_pdf_note(path: Path) -> str:
    """PDFはテキスト層がないスキャンが多い。Haikuに画像として渡すのは
    コスト増のため、ここではプレースホルダを返しユーザーに文字起こし済み
    テキストの用意を促す。"""
    return f"[PDF: {path.name} — スキャンPDFはテキスト化して .txt を同名で置いてください]"


def scan_folder(folder: Path) -> dict:
    """フォルダ内の教材を {doc_id: {path, text}} に集める。"""
    docs = {}
    for f in sorted(folder.iterdir()):
        if f.name.startswith("~$") or f.name.startswith("."):
            continue
        doc_id = re.sub(r"[^A-Za-z0-9_]+", "_", f.stem).strip("_").lower()
        if f.suffix.lower() == ".docx":
            docs[doc_id] = {"path": str(f), "text": extract_docx(f)}
        elif f.suffix.lower() in (".txt", ".md"):
            docs[doc_id] = {"path": str(f), "text": f.read_text(encoding="utf-8")}
        elif f.suffix.lower() == ".pdf":
            txt = f.with_suffix(".txt")
            if txt.exists():
                docs[doc_id] = {"path": str(f), "text": txt.read_text(encoding="utf-8")}
            else:
                docs[doc_id] = {"path": str(f), "text": extract_pdf_note(f)}
    return docs


def save_index(docs: dict, out: Path):
    out.write_text(json.dumps(docs, ensure_ascii=False, indent=1), encoding="utf-8")
