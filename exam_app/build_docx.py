"""exam_draft.json → Word 3点セット（問題用紙・解答用紙・模範解答）。

体裁は最小限（A4・明朝/Times）。細かい整形は既存テンプレート運用
（scripts/make_answersheet_comm2.py 方式）を推奨し、これは叩き台。
skills/exam-answersheet: 解答用紙と模範解答は同一関数から生成し二重管理を避ける。
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _para(doc, text, size=10.5, bold=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "MS Mincho")
    return p


def build_exam(draft: dict, outdir: Path):
    doc = Document()
    _para(doc, draft["exam"]["title"], size=12, bold=True, center=True)
    for sec in draft["sections"]:
        pts = f"（{sec['points_each']}点×{len(sec['questions'])}＝" \
              f"{sec['points_each'] * len(sec['questions'])}点）"
        _para(doc, f"\n{sec['no']}　{sec.get('instructions', '').strip()} {pts}",
              bold=True)
        for q in sec["questions"]:
            _para(doc, f"({q['number']})　{q['body']}")
    out = outdir / "exam_draft.docx"
    doc.save(str(out))
    return out


def build_answersheet(draft: dict, outdir: Path, model: bool):
    """model=False: 解答用紙（空欄） / True: 模範解答（太字で記入）"""
    doc = Document()
    suffix = "模範解答" if model else "解答用紙"
    _para(doc, f"{draft['exam']['title']}　{suffix}", size=12, bold=True, center=True)
    for sec in draft["sections"]:
        _para(doc, f"\n{sec['no']}　【{sec['points_each']}点×{len(sec['questions'])}】",
              bold=True)
        t = doc.add_table(rows=len(sec["questions"]), cols=2)
        t.style = "Table Grid"
        for i, q in enumerate(sec["questions"]):
            t.rows[i].cells[0].text = str(q["number"])
            t.rows[i].cells[0].width = Pt(30)
            if model:
                cell = t.rows[i].cells[1]
                cell.text = ""
                run = cell.paragraphs[0].add_run(q["answer"])
                run.font.bold = True
                run.font.size = Pt(10)
    name = "modelanswer_draft.docx" if model else "answersheet_draft.docx"
    out = outdir / name
    doc.save(str(out))
    return out


def build_all(draft: dict, outdir: Path) -> list:
    return [
        build_exam(draft, outdir),
        build_answersheet(draft, outdir, model=False),
        build_answersheet(draft, outdir, model=True),
    ]
