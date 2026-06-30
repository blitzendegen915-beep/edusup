from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = '游明朝'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

def add_run(para, text, size=11, bold=False, color=None, highlight=None):
    r = para.add_run(text)
    r.font.name = '游明朝'
    r.font.size = Pt(size)
    r.bold = bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    if color:
        r.font.color.rgb = RGBColor(*color)
    if highlight:
        rPr = r._element.get_or_add_rPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), highlight)
        rPr.append(shd)
    return r

def para(text='', align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=False, space_before=0, space_after=0, color=None, highlight=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        add_run(p, text, size=size, bold=bold, color=color, highlight=highlight)
    return p

YELLOW = 'FFFF00'

# ── 宛名・日付・差出人 ──
p = para('ラグビー部保護者各位', size=11, space_after=0)

p_date = para(align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
add_run(p_date, '2026年　　月　　日', size=11, highlight=YELLOW)

para('駒澤大学高等学校', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_after=0)

p_principal = para(align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
add_run(p_principal, '校長　井上誠二', size=11)

para('ラグビー部顧問　占部涼也　畠山和真', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_after=6)

# ── タイトル ──
para('ラグビー部　夏合宿のご案内', align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_before=6, space_after=8)

# ── 挨拶文 ──
greet = doc.add_paragraph()
greet.paragraph_format.first_line_indent = Pt(11)
greet.paragraph_format.space_before = Pt(0)
greet.paragraph_format.space_after  = Pt(4)
add_run(greet,
    '保護者各位におかれましては、平素より本校の教育活動にご理解とご協力を賜り、'
    '心より御礼申し上げます。まもなく夏休みとなりますが、夏合宿のご案内をさせて頂きます。'
    'ラグビー部においては、今年もメンバーの強化とチームの団結力を高めるため、'
    '夏合宿を下記の要領にて実施致します。奮ってご参加下さいますよう、よろしくお願い致します。',
    size=10.5)

greet2 = doc.add_paragraph()
greet2.paragraph_format.first_line_indent = Pt(11)
greet2.paragraph_format.space_before = Pt(0)
greet2.paragraph_format.space_after  = Pt(4)
add_run(greet2,
    'なお引率指導コーチの宿泊費・交通費に関しましては、部員負担となります。'
    '何卒ご理解とご了承を下さいますよう、合わせてよろしくお願い申し上げます。',
    size=10.5)

greet3 = doc.add_paragraph()
greet3.paragraph_format.first_line_indent = Pt(11)
greet3.paragraph_format.space_before = Pt(0)
greet3.paragraph_format.space_after  = Pt(4)
add_run(greet3,
    'また、今年の夏は例年よりも更に暑くなることが想定されます。'
    '熱中症や怪我にはスタッフ一同気を付けますが、万が一の場合に備えて、'
    'こちらからの連絡が取れるよう、よろしくお願い致します。',
    size=10.5)

greet4 = doc.add_paragraph()
greet4.paragraph_format.first_line_indent = Pt(11)
greet4.paragraph_format.space_before = Pt(0)
greet4.paragraph_format.space_after  = Pt(6)
add_run(greet4,
    '体調不良やケガなどで急遽不参加になる場合は、可能な限り返金しますが、'
    '交通費やグラウンド使用料など人数割の部分に関しては返金致しかねますのでご了承ください。',
    size=10.5)

# ── 記 ──
para('記', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_before=2, space_after=6)

# ── 期間 ──
p_period = doc.add_paragraph()
p_period.paragraph_format.space_after = Pt(2)
add_run(p_period, '期間　', size=11, bold=True)
add_run(p_period, '８月１日（土）～８月７日（金）　６泊７日', size=11)

# ── 日程 ──
p_sched = doc.add_paragraph()
p_sched.paragraph_format.space_after = Pt(1)
add_run(p_sched, '日程　', size=11, bold=True)
add_run(p_sched, '８月１日（土）午前　移動（高校7時半集合）　／　午後　練習', size=10.5)

schedule_lines = [
    ('８月２日（日）', '午前　　　　　　　　　　　　／　午後　'),
    ('８月３日（月）', '午前　　　　　　　　　　　　／　午後　'),
    ('８月４日（火）', '午前　　　　　　　　　　　　／　午後　'),
    ('８月５日（水）', '午前　　　　　　　　　　　　／　午後　'),
    ('８月６日（木）', '午前　　　　　　　　　　　　／　午後　帰京'),
]
# Note: 8/7 is departure day (last day checkout + travel back)
# Actually 8/1-8/7, 8/7 afternoon is 帰京
# Let me fix: 8/1 arrive, 8/2-6 practice/games, 8/7 morning practice + afternoon 帰京

schedule_data = [
    ('８月２日（日）', '対戦校未定'),
    ('８月３日（月）', '対戦校未定'),
    ('８月４日（火）', '対戦校未定'),
    ('８月５日（水）', '対戦校未定'),
    ('８月６日（木）', '対戦校未定'),
    ('８月７日（金）', '午前　練習　／　午後　帰京'),
]
for date, content in schedule_data:
    ps = doc.add_paragraph()
    ps.paragraph_format.space_before = Pt(0)
    ps.paragraph_format.space_after  = Pt(1)
    ps.paragraph_format.left_indent = Pt(42)
    add_run(ps, f'{date}', size=10.5)
    if '未定' in content:
        add_run(ps, f'　{content}', size=10.5, highlight=YELLOW)
    else:
        add_run(ps, f'　{content}', size=10.5)

p_note_ig = doc.add_paragraph()
p_note_ig.paragraph_format.space_before = Pt(2)
p_note_ig.paragraph_format.space_after = Pt(4)
add_run(p_note_ig, '※グラウンドNo.については、Instagramの投稿をご確認ください。', size=10)

# ── 宿泊 ──
p_hotel = doc.add_paragraph()
p_hotel.paragraph_format.space_after = Pt(2)
add_run(p_hotel, '宿泊　', size=11, bold=True)
add_run(p_hotel, '菅平ホテル：〒386-2204　長野県上田市菅平高原1262-4　　電話：0268-74-2001', size=10.5)

# ── 持物 ──
p_items = doc.add_paragraph()
p_items.paragraph_format.space_after = Pt(2)
add_run(p_items, '持物　', size=11, bold=True)
add_run(p_items, '保険証、部活着等（詳しくはマネージャー経由でお伝えします）', size=10.5)

# ── 費用 ──
p_cost_header = doc.add_paragraph()
p_cost_header.paragraph_format.space_before = Pt(4)
p_cost_header.paragraph_format.space_after = Pt(2)
add_run(p_cost_header, '費用', size=11, bold=True)

p_cost_note = doc.add_paragraph()
p_cost_note.paragraph_format.space_after = Pt(4)
p_cost_note.paragraph_format.left_indent = Pt(14)
add_run(p_cost_note,
    '※ 近年の物価高の影響により、宿泊代・食事代等が値上がりしております。'
    '保護者の皆様にはご負担をおかけしてしまいますが、何卒ご理解いただきたくお願い申し上げます。',
    size=10)

cost_items = [
    ('・宿泊費（1泊3食×6泊）', '￥55,020', '（￥9,170×6泊分）', None),
    ('・増昼食代（最終日1日分）', '￥1,100', '', None),
    ('・BBQ代', '￥1,100', '', None),
    ('・グラウンド使用料', '￥2,152', '（グラウンド使用料合計￥99,000を46人で割る）', None),
    ('・交通費', '￥9,678', '（バス往復手配￥445,200を46人で割る）', None),
    ('・雑費', '￥6,000', '（雑費1人1日1,000円×6日分）', YELLOW),
    ('・コーチ謝礼・宿泊・交通費', '未定', '（選手のみ負担）', YELLOW),
    ('・予備費', '未定', '', YELLOW),
]

for label, amount, note, hl in cost_items:
    p_c = doc.add_paragraph()
    p_c.paragraph_format.space_before = Pt(0)
    p_c.paragraph_format.space_after  = Pt(1)
    p_c.paragraph_format.left_indent = Pt(14)
    add_run(p_c, label, size=10.5)
    add_run(p_c, f'　{amount}', size=10.5, bold=True, highlight=hl)
    if note:
        add_run(p_c, f'　{note}', size=10, highlight=hl)

# ── 合計 ──
p_total = doc.add_paragraph()
p_total.paragraph_format.space_before = Pt(6)
p_total.paragraph_format.space_after = Pt(2)
p_total.paragraph_format.left_indent = Pt(14)
add_run(p_total, '★合宿費合計　', size=11, bold=True)
add_run(p_total, '￥75,050＋コーチ謝礼等＋予備費（確定後ご案内します）', size=11, bold=True, highlight=YELLOW)

p_total_note = doc.add_paragraph()
p_total_note.paragraph_format.space_after = Pt(6)
p_total_note.paragraph_format.left_indent = Pt(14)
add_run(p_total_note, '※確定分のみの小計：￥75,050（宿泊費＋増昼食＋BBQ＋グラウンド＋交通費＋雑費）', size=10)

# ── 振込先 ──
p_bank_header = doc.add_paragraph()
p_bank_header.paragraph_format.space_before = Pt(2)
p_bank_header.paragraph_format.space_after = Pt(2)
add_run(p_bank_header, '振込　', size=11, bold=True)
add_run(p_bank_header, '三菱UFJ銀行　世田谷支店　普通預金', size=10.5)

p_bank_detail = doc.add_paragraph()
p_bank_detail.paragraph_format.space_after = Pt(2)
p_bank_detail.paragraph_format.left_indent = Pt(42)
add_run(p_bank_detail, '店番130　口座番号0887504　駒澤大学高校　ラグビー部　畠山和真', size=10.5)

p_bank_note = doc.add_paragraph()
p_bank_note.paragraph_format.space_after = Pt(4)
p_bank_note.paragraph_format.left_indent = Pt(42)
add_run(p_bank_note, '※お振込みの際には、ご依頼人名の欄に「学年（生徒氏名）」のご入力をお願いいたします。', size=10)

# ── 締切 ──
p_deadline = doc.add_paragraph()
p_deadline.paragraph_format.space_after = Pt(4)
add_run(p_deadline, '締切　', size=11, bold=True)
add_run(p_deadline, '　月　　日（　　）', size=11, highlight=YELLOW)

# ── 引率 ──
p_escort = doc.add_paragraph()
p_escort.paragraph_format.space_after = Pt(2)
add_run(p_escort, '引率顧問', size=11, bold=True)
add_run(p_escort, '　占部涼也・畠山和真', size=10.5)

p_coach = doc.add_paragraph()
p_coach.paragraph_format.space_after = Pt(2)
add_run(p_coach, '外部コーチ', size=11, bold=True)
add_run(p_coach, '　　名', size=10.5, highlight=YELLOW)

p_student = doc.add_paragraph()
p_student.paragraph_format.space_after = Pt(2)
add_run(p_student, '参加生徒', size=11, bold=True)
add_run(p_student, '　42名（選手38名・マネージャー4名）', size=10.5)

p_total_people = doc.add_paragraph()
p_total_people.paragraph_format.space_after = Pt(6)
add_run(p_total_people, '計　', size=11, bold=True)
add_run(p_total_people, '　名', size=11, highlight=YELLOW)

# ── 以上 ──
para('以　上', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_before=4)

os.makedirs('/home/user/teacher-automation-lab/output', exist_ok=True)
out = '/home/user/teacher-automation-lab/output/2026_夏合宿のご案内.docx'
doc.save(out)
print(f'完了: {out}')
