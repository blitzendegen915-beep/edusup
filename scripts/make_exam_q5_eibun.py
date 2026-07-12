"""
英語コミュニケーションII 1学期期末試験
大問5 英文解釈問題 単独Wordファイル出力
入門英文問題精講 問題番号 33,34,37,38,41,42 を使用
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, name_ja="MS Mincho", name_en="Times New Roman",
             size=10.5, bold=False, italic=False, underline=False):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_ja)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def highlight_yellow(run):
    r = run._r
    rPr = r.get_or_add_rPr()
    hl = OxmlElement('w:highlight')
    hl.set(qn('w:val'), 'yellow')
    rPr.append(hl)

def add_run(para, text, bold=False, italic=False, underline=False,
            size=10.5, yellow=False):
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, underline=underline)
    if yellow:
        highlight_yellow(run)
    return run

def add_boxed_number(para, number):
    run = para.add_run(str(number))
    set_font(run, size=11, bold=True)
    r = run._r
    rPr = r.get_or_add_rPr()
    bdr = OxmlElement('w:bdr')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '8')
    bdr.set(qn('w:space'), '2')
    bdr.set(qn('w:color'), 'auto')
    rPr.append(bdr)
    return run

def set_para(para, align=None, sb=0, sa=0, il=0, fi=0, ls=None):
    pf = para.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    pf.left_indent = Cm(il)
    pf.first_line_indent = Cm(fi)
    if align:
        para.alignment = align
    if ls:
        pf.line_spacing = Pt(ls)

def set_page_layout(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def build_q5():
    doc = Document()
    set_page_layout(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

    # ======================================================
    # タイトル
    # ======================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(title, sb=0, sa=4)
    add_run(title, "２０２６年　２学年受験コース　英語コミュニケーションⅡ　１学期期末試験", bold=True, size=12)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(sub, sb=0, sa=8)
    add_run(sub, "大問５　英文解釈", bold=True, size=11)

    # ======================================================
    # 大問5ヘッダー
    # ======================================================
    header = doc.add_paragraph()
    set_para(header, sb=4, sa=2)
    add_boxed_number(header, 5)
    add_run(header, "　以下の英文解釈に関する問題に答えなさい。", bold=True)

    # ======================================================
    # Part 1 指示文
    # ======================================================
    inst1 = doc.add_paragraph()
    set_para(inst1, sb=4, sa=2, il=0.3)
    add_run(inst1,
        "下の例にならい、次の文１〜４に SVOCM の記号を振り構造を表しなさい。また、それにもとづいた英文の"
        "日本語訳を書きなさい。ただし、1 つの単語につき 2 つ以上の記号・下線は付さなくて良い。"
        "文全体の大きな構造がわかるように図示すること。")

    # 例
    ex_label = doc.add_paragraph()
    set_para(ex_label, sb=2, sa=0, il=0.3)
    add_run(ex_label, "例", bold=True, size=9)

    ex_key = doc.add_paragraph()
    set_para(ex_key, sb=0, sa=2, il=0.5)
    add_run(ex_key,
        "（記号）S → 四角　、V → 下線　、O → 二重下線　、C → 波線　、M → （まるカッコ）", size=9)

    ex_sent = doc.add_paragraph()
    set_para(ex_sent, sb=0, sa=6, il=0.5)
    add_run(ex_sent,
        "Developing nations were under-represented, (with competitors facing obstacles "
        "ranging from prejudice against disability to the high cost of wheelchairs).", size=9.5)

    # ======================================================
    # Part 1 問1〜4（SVOCM + 日本語訳）
    # ======================================================

    problems_p1 = [
        (
            "She also saw the poverty of her people and the hard lives of so many women "
            "who were fighting against such basic problems as lack of food, firewood and water, "
            "and against unemployment.",
            # 解答
            "S[She] V[saw] O[the poverty of her people] and O[the hard lives of so many women "
            "[who were fighting against (such basic problems as lack of food, firewood and water), "
            "and against unemployment]].",
            # 日本語訳
            "彼女はまた、自分の民の貧困と、食料・薪・水の不足や失業といった基本的な問題と闘う数多くの女性たちの"
            "苦しい生活を目の当たりにした。",
            # 補足問 なし
            None, None
        ),
        (
            "In March, the AI-based computer program AlphaGo, developed by DeepMind, shocked the world "
            "when it defeated South Korean go grandmaster Lee Sedol in a five-game match of the ancient "
            "board game that requires deep insight.",
            # 解答
            "M(In March), S[the AI-based computer program AlphaGo, (developed by DeepMind),] "
            "V[shocked] O[the world] M(when it defeated South Korean go grandmaster Lee Sedol "
            "in a five-game match of the ancient board game [that requires deep insight]).",
            # 日本語訳
            "3月、DeepMindが開発したAIベースのコンピュータプログラム AlphaGo は、深い洞察力を必要とする"
            "この古代ボードゲームの5番勝負で韓国の囲碁の名人、李世乭を破り、世界に衝撃を与えた。",
            # 補足問
            "下線部が示す内容を日本語で具体的に答えなさい。",
            "AlphaGoが人間のプロ棋士（李世乭）との5番勝負に勝利したこと。"
        ),
        (
            "Although we know that the earth revolves around the sun, we cannot recite "
            "the astronomical observations and calculations that led to that conclusion.",
            # 解答
            "M(Although S[we] V[know] O〈that the earth revolves around the sun〉), "
            "S[we] V[cannot recite] O[the astronomical observations and calculations "
            "[that led to that conclusion]].",
            # 日本語訳
            "地球が太陽の周りを公転していることは知っていても、私たちはその結論に至った天文学的な観測や"
            "計算を暗唱することはできない。",
            None, None
        ),
        (
            "Psychologists who believe that willpower is a limited resource say using up "
            "our willpower is the main reason that some of us fail to achieve our goals.",
            # 解答
            "S[Psychologists [who believe 〈that willpower is a limited resource〉]] "
            "V[say] O〈using up our willpower is the main reason [that some of us fail to achieve our goals]〉.",
            # 日本語訳
            "意志力は限られた資源だと考える心理学者たちは、意志力を使い果たすことが、"
            "私たちの一部が目標を達成できない主な理由だと言う。",
            None, None
        ),
    ]

    for i, (sentence, answer, translation, sub_q, sub_ans) in enumerate(problems_p1, 1):
        p_num = doc.add_paragraph()
        set_para(p_num, sb=6, sa=0, il=0.3)
        add_run(p_num, f"　{i}．", bold=True)
        add_run(p_num, sentence)

        # 記号欄
        label_s = doc.add_paragraph()
        set_para(label_s, sb=1, sa=0, il=0.5)
        add_run(label_s, "（記号）", size=9.5)

        blank_s = doc.add_paragraph()
        set_para(blank_s, sb=0, sa=0, il=0.5)
        add_run(blank_s, answer, yellow=True, size=9.5)

        # 日本語訳欄
        label_t = doc.add_paragraph()
        set_para(label_t, sb=2, sa=0, il=0.5)
        add_run(label_t, "（訳）", size=9.5)

        blank_t = doc.add_paragraph()
        set_para(blank_t, sb=0, sa=0, il=0.5)
        add_run(blank_t, translation, yellow=True, size=9.5)

        # 補足問
        if sub_q:
            sq_p = doc.add_paragraph()
            set_para(sq_p, sb=2, sa=0, il=0.5)
            add_run(sq_p, "（問）", bold=True, size=9.5)
            add_run(sq_p, sub_q, size=9.5)
            sa_p = doc.add_paragraph()
            set_para(sa_p, sb=0, sa=0, il=0.7)
            add_run(sa_p, sub_ans, yellow=True, size=9.5)

        doc.add_paragraph()

    # ======================================================
    # Part 2 指示文（正誤判定）
    # ======================================================
    inst2 = doc.add_paragraph()
    set_para(inst2, sb=8, sa=2, il=0.3)
    add_run(inst2,
        "下の５〜８の構文解釈が正しければア、間違っていればイを書きなさい。"
        "５〜６は SVOCM について、７〜８は句・節について（記号は"
        "〈名詞句・節〉　［形容詞句・節］　（副詞句・節））である。")

    problems_p2 = [
        (
            # 番号5: X-rays / SVOC問題
            "5.",
            "S[X-rays] V[allowed] O[them] C[to look into their patients, "
            "identify where there were problems, and cure them].",
            "ア",
            "ア（正しい）：allow O to do = SVOC の構造。O=them, C=to look...となる。"
        ),
        (
            # 番号6: Willpower / 間違い（a mysteriousのみをCとしている）
            "6.",
            "S[Willpower] V[is] C[a mysterious] M（force [that helps us control "
            "our actions and achieve our goals]）.",
            "イ",
            "イ（間違い）：C は「a mysterious force [that...]」全体。"
            "「a mysterious」だけをCとし「force...」をMと分けるのは誤り。"
        ),
        (
            # 番号7: Chopik / 正しい
            "7.",
            "S[Chopik] V[says] O〈he isn't suggesting 〈we ignore our families〉, "
            "but 〈that friends make us feel better〉〉.",
            "ア",
            "ア（正しい）：says の目的語節（名詞節）の中に、suggesting の目的語節が2つ並列されている。"
        ),
        (
            # 番号8: With friends / 間違い（they provide...をMとしている）
            "8.",
            "M（With friends）S[you] V[are] C[more likely to do activities] "
            "M（— they provide an outlet）.",
            "イ",
            "イ（間違い）：ダッシュ以降の「they provide an outlet」は独立した節であり、"
            "修飾語句Mではない。"
        ),
    ]

    sentences_p2 = [
        "X-rays allowed them to look into their patients, identify where there were problems, and cure them.",
        "Willpower is a mysterious force that helps us control our actions and achieve our goals.",
        "Chopik says he isn't suggesting we ignore our families, but that friends make us feel better.",
        "With friends you are more likely to do activities — they provide an outlet.",
    ]

    for (num, analysis, ans, explanation), sentence in zip(problems_p2, sentences_p2):
        p_sent = doc.add_paragraph()
        set_para(p_sent, sb=4, sa=0, il=0.3)
        add_run(p_sent, f"　{num}　", bold=True)
        add_run(p_sent, sentence)

        p_ana = doc.add_paragraph()
        set_para(p_ana, sb=1, sa=0, il=0.7)
        add_run(p_ana, analysis, size=9.5)

        p_ans = doc.add_paragraph()
        set_para(p_ans, sb=1, sa=4, il=0.7)
        add_run(p_ans, ans, bold=True, yellow=True)
        add_run(p_ans, "　" + explanation, size=9, yellow=True)

    # 答え一覧
    ans_row = doc.add_paragraph()
    set_para(ans_row, sb=6, sa=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(ans_row, "５ア　６イ　７ア　８イ", bold=True, yellow=True)

    # ======================================================
    # 保存
    # ======================================================
    import os
    out = "/home/user/teacher-automation-lab/output/exam_comm2_q5_eibun.docx"
    os.makedirs("/home/user/teacher-automation-lab/output", exist_ok=True)
    doc.save(out)
    print(f"保存完了: {out}")

if __name__ == "__main__":
    build_q5()
