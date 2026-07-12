"""
英語コミュニケーションII 1学期期末試験
解答用紙・模範解答 生成スクリプト
"""

import os
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "/home/user/teacher-automation-lab/materials/answersheet_comm2_template_new.docx"

# ========== 模範解答データ ==========

ANSWERS = {
    # 大問1: 1点×9（空所補充）
    'q1': [
        "commute", "universal", "must have left",
        "greenhouse", "nuance", "conduct",
        "goes without", "facial", "safely",
    ],

    # 大問2: 2点×3
    # Q1=不要文番号, Q2=空欄Aの記号, Q3=内容一致の記号
    'q2': {'1': '③', '2': 'ア', '3': 'ア'},

    # 大問3: 1点×4（空所補充）
    'q3': {'1': 'ウ', '2': 'エ', '3': 'ア', '4': 'イ'},

    # 大問4: 2点×5（語句挿入 → 前後の単語を記入）
    # 挿入語句リスト（解答用紙に印刷する）
    'q4_words': ["so that", "such as", "quicker", "from", "well"],
    # 模範解答: (挿入箇所の直前の単語, 直後の単語)
    'q4_answers': [
        ("systems",       "they"),          # 1. so that
        ("infrastructure","the"),           # 2. such as
        ("now",           "to"),            # 3. quicker
        ("vehicles",      "entering"),      # 4. from
        ("may",           "Copenhagenize"), # 5. well
    ],

    # 大問5: 英文解釈（構文+訳 ×2、Part2 ×4）
    'q5_trans': [
        "大人になることは段階的なプロセスであり、若者たちがようやく若い大人として扱われるほど賢くなったまさにその時こそ、彼らにお酒を自由に飲ませる時ではない。",
        "地球が太陽の周りを公転していることは知っていても、私たちはその結論に至った天文学的な観測や計算を暗唱することはできない。",
    ],
    'q5_part2': {"3": "ア", "4": "ア", "5": "イ", "6": "イ"},

    # 大問6: 2点×5（並べ替え → 2番目と5番目）
    # (2番目の語句, 5番目の語句)
    # 問1: Those who exercise regularly feel better about themselves.
    # 問2: Don't put off until tomorrow what you can do today.
    # 問3: She is the girl who I think will win the prize.
    # 問4: What you need is to get enough sleep. ※別解検討中
    # 問5: It is important to make friends with people with whom you can share new experiences.
    'q6': [
        ('who',   'better'),   # 問1
        ('off',   'what'),     # 問2
        ('the',   'I'),        # 問3
        ('you',   'enough'),   # 問4（差し替え未確定）
        ('make',  'people'),   # 問5
    ],

    # 大問7: 3点×5（内容一致）
    'q7': ["3", "4", "2", "3", "3"],

    # 大問8: 1点×10（語群選択）
    'q8': ["サ", "キ", "エ", "ケ", "コ", "オ", "ア", "ウ", "ク", "イ"],
}

# ========== ヘルパー ==========

def set_font(run, size=10, bold=False, name_ja="MS Mincho", name_en="Times New Roman"):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_ja)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def write_cell(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """セルの内容を書き換える（セルの罫線・背景は変更しない）"""
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _clear_and_set_tc(tc, text, size=9, bold=False, align='center'):
    """XML レベルでセルの段落を書き換える（rebuild_q4 用）"""
    for p in list(tc.findall(qn('w:p'))):
        tc.remove(p)

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size * 2))
        rPr.append(szCs)
        if bold:
            rPr.append(OxmlElement('w:b'))
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        p.append(r)

    tc.append(p)


def rebuild_q4(doc, words, answers, model=False):
    """
    大問4 (Table 4) を 5行×4列 に作り直す。
    列構成: [番号] [前の単語] [挿入語句] [後の単語]
    """
    t4 = doc.tables[4]
    tbl = t4._tbl

    # 参照セルを取得（罫線・幅のスタイルを流用）
    ref_tr = tbl.findall(qn('w:tr'))[0]
    ref_tcs = ref_tr.findall(qn('w:tc'))
    ref_num   = deepcopy(ref_tcs[0])   # 番号セルのスタイル
    ref_blank = deepcopy(ref_tcs[1])   # 空欄セルのスタイル

    # 既存の行をすべて削除
    for tr in list(tbl.findall(qn('w:tr'))):
        tbl.remove(tr)

    # 5問分の行を追加
    for i, (word, (before, after)) in enumerate(zip(words, answers)):
        tr = OxmlElement('w:tr')

        tc_num = deepcopy(ref_num)
        _clear_and_set_tc(tc_num, str(i + 1), size=9, bold=False)
        tr.append(tc_num)

        tc_before = deepcopy(ref_blank)
        _clear_and_set_tc(tc_before, before if model else '', size=9, bold=model)
        tr.append(tc_before)

        tc_word = deepcopy(ref_blank)
        _clear_and_set_tc(tc_word, word, size=9, bold=False)
        tr.append(tc_word)

        tc_after = deepcopy(ref_blank)
        _clear_and_set_tc(tc_after, after if model else '', size=9, bold=model)
        tr.append(tc_after)

        tbl.append(tr)


# ========== メイン ==========

def build(model=False):
    doc = Document(TEMPLATE)

    # --------------------------------------------------
    # T1: 大問1（9問）
    # 各行: [番号, blank×4, 番号, blank×2]
    # 左問: col0=番号, col1=解答　右問: col5=番号, col6=解答
    # --------------------------------------------------
    t1 = doc.tables[1]
    ans = ANSWERS['q1'] if model else [''] * 9
    pairs = [(0,1),(2,3),(4,5),(6,7),(8,None)]
    for ri, (li, ri2) in enumerate(pairs):
        row = t1.rows[ri]
        write_cell(row.cells[1], ans[li], size=9, bold=model)
        if ri2 is not None:
            write_cell(row.cells[6], ans[ri2], size=9, bold=model)

    # --------------------------------------------------
    # T2: 大問2（Q1=不要文, Q2=空欄A, Q3=内容一致）
    # セル: [２, 1, blank, 2A→2, blank, 2B→3, blank]
    # --------------------------------------------------
    t2 = doc.tables[2]
    row2 = t2.rows[0]
    # ラベルを 2A→2, 2B→3 に変更
    write_cell(row2.cells[3], '2', size=9, bold=False)
    write_cell(row2.cells[5], '3', size=9, bold=False)
    # 解答欄
    write_cell(row2.cells[2], ANSWERS['q2']['1'] if model else '', size=11, bold=model)
    write_cell(row2.cells[4], ANSWERS['q2']['2'] if model else '', size=11, bold=model)
    write_cell(row2.cells[6], ANSWERS['q2']['3'] if model else '', size=11, bold=model)

    # --------------------------------------------------
    # T3: 大問3（1〜4）
    # [３, 1, blank, 2, blank, 3, blank, 4, blank]
    # --------------------------------------------------
    t3 = doc.tables[3]
    row3 = t3.rows[0]
    for i, col in enumerate([2, 4, 6, 8]):
        write_cell(row3.cells[col], ANSWERS['q3'][str(i+1)] if model else '', size=11, bold=model)

    # --------------------------------------------------
    # T4: 大問4（5問×2ボックス → 前後の単語）
    # 作り直し: [番号][前の単語][挿入語句][後の単語] × 5行
    # --------------------------------------------------
    rebuild_q4(
        doc,
        ANSWERS['q4_words'],
        ANSWERS['q4_answers'] if model else [('', '')] * 5,
        model=model,
    )

    # --------------------------------------------------
    # T5: 大問5（英文解釈）
    # row0: 英文1（変更しない）  row1: 訳1の空欄
    # row2: 英文2（変更しない）  row3: 訳2の空欄
    # row4: Part2（3〜6の解答）
    # --------------------------------------------------
    t5 = doc.tables[5]
    trans = ANSWERS['q5_trans'] if model else ['', '']
    p2    = ANSWERS['q5_part2'] if model else {str(k): '' for k in range(3, 7)}

    write_cell(t5.rows[1].cells[1], trans[0], size=8, bold=model, align=WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(t5.rows[3].cells[1], trans[1], size=8, bold=model, align=WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(t5.rows[4].cells[2], p2['3'], size=10, bold=model)
    write_cell(t5.rows[4].cells[4], p2['4'], size=10, bold=model)
    write_cell(t5.rows[4].cells[6], p2['5'], size=10, bold=model)
    write_cell(t5.rows[4].cells[8], p2['6'], size=10, bold=model)

    # --------------------------------------------------
    # T6: 大問6（並べ替え 2番目/5番目 × 5問）
    # [６, 1, blank, 2, blank, 3, blank, 4, blank, 5, blank]
    # --------------------------------------------------
    t6 = doc.tables[6]
    ans6 = ANSWERS['q6'] if model else [('', '')] * 5
    for i, col in enumerate([2, 4, 6, 8, 10]):
        a2, a5 = ans6[i]
        write_cell(t6.rows[0].cells[col],
                   f'②{a2}　⑤{a5}' if model else '', size=8, bold=model)

    # --------------------------------------------------
    # T7: 大問7（3点×5）
    # [７, 1, blank, 2, blank, 3, blank, 4, blank, 5, blank]
    # --------------------------------------------------
    t7 = doc.tables[7]
    ans7 = ANSWERS['q7'] if model else [''] * 5
    for i, col in enumerate([2, 4, 6, 8, 10]):
        write_cell(t7.rows[0].cells[col], ans7[i] if model else '', size=11, bold=model)

    # --------------------------------------------------
    # T8: 大問8（10問、2行）
    # row0: col1=q1, col3=q2, col5=q3, col7=q4, col9=q5
    # row1: col1=q6, col3=q7, col5=q8, col7=q9, col9=q10
    # --------------------------------------------------
    t8 = doc.tables[8]
    ans8 = ANSWERS['q8'] if model else [''] * 10
    for i, col in enumerate([1, 3, 5, 7, 9]):
        write_cell(t8.rows[0].cells[col], ans8[i],     size=10, bold=model)
        write_cell(t8.rows[1].cells[col], ans8[i + 5], size=10, bold=model)

    # --------------------------------------------------
    # 保存
    # --------------------------------------------------
    os.makedirs("/home/user/teacher-automation-lab/output", exist_ok=True)
    suffix = "modelanswer" if model else "answersheet"
    out = f"/home/user/teacher-automation-lab/output/{suffix}_comm2_final.docx"
    doc.save(out)
    print(f"保存完了: {out}")


if __name__ == "__main__":
    build(model=False)
    build(model=True)
