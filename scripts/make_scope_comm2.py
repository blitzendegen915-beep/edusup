"""
英語コミュニケーションII 1学期期末試験 出題範囲 Word生成
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, name_ja="MS Mincho", name_en="Times New Roman",
             size=10.5, bold=False, underline=False):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_ja)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def add_run(para, text, bold=False, underline=False, size=10.5):
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, underline=underline)
    return run

def sp(para, sb=0, sa=0, il=0):
    pf = para.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    pf.left_indent = Cm(il)

def add_box_text(para, text, size=10.5):
    """枠線付きテキスト（出題テキスト・配点ラベル）"""
    run = para.add_run(text)
    set_font(run, size=size, underline=True)
    return run

def set_page_layout(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

def add_word_table(doc, words, answers):
    """語句挿入問題の語句ボックス"""
    table = doc.add_table(rows=1, cols=len(words))
    table.style = 'Table Grid'
    for i, word in enumerate(words):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(word)
        set_font(run, size=10, bold=True if i == 0 else False)
    return table

def build_scope():
    doc = Document()
    set_page_layout(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

    # ======================================================
    # タイトル
    # ======================================================
    title = doc.add_paragraph()
    sp(title, sb=0, sa=6)
    add_run(title, "２年受験コース　コミュⅡ　１学期期末試験　出題範囲", bold=True, size=12)

    # ======================================================
    # 出題テキスト
    # ======================================================
    label1 = doc.add_paragraph()
    sp(label1, sb=4, sa=2)
    add_box_text(label1, "出題テキスト")

    items_text = [
        ("・Heartening Ⅱ L3,4：　教科書本文/ワークブック", False),
        ("・入門英文問題精講（入門）：", False),
        ("　　授業で扱ったもの（問題番号　25,26,29,30,33,34,37,38,41,42）", False),
        ("　　※加えて 27,28,31,32,35 からも出題します。", False),
        ("・Neo現代 Unit2", False),
    ]
    for text, bold in items_text:
        p = doc.add_paragraph()
        sp(p, sb=0, sa=0)
        add_run(p, text, bold=bold)

    # ======================================================
    # 出題内容・配点
    # ======================================================
    label2 = doc.add_paragraph()
    sp(label2, sb=10, sa=2)
    add_box_text(label2, "出題内容・配点（予定）")

    total_p = doc.add_paragraph()
    sp(total_p, sb=0, sa=2)
    add_run(total_p, "テストは８０点満点（＋Speaking １０点＋Listening １０点　＝　１００点）")

    listen_p = doc.add_paragraph()
    sp(listen_p, sb=0, sa=4)
    add_run(listen_p, "※リスニングは６月２９日（月）に実施します。")

    scores = [
        ("ワークブック　　　　１０点：", "　穴埋め問題"),
        ("HearteningⅡ　　　　２０点：", "　※どこにあったの問題など"),
        ("入門　　　　　　　　１５点：", "　与えられた文に記号を振り構造を示す。それに基づいた日本語訳をする。\n"
                                         "　　　　　　　　　　　　　（＜名詞句・節＞, [形容詞句・節], (副詞句・節)。SVOCM）"),
        ("動画　　　　　　　　１０点：", "　並び替え問題"),
        ("Neo現代 Unit2　　　１０点：", "　本文の穴埋め問題"),
        ("初見問題　　　　　　１５点：", "　英検２級レベル　リーディング"),
    ]

    for label, content in scores:
        p = doc.add_paragraph()
        sp(p, sb=0, sa=0)
        add_run(p, label, bold=True)
        add_run(p, content)

    # ======================================================
    # 勉強の仕方
    # ======================================================
    label3 = doc.add_paragraph()
    sp(label3, sb=10, sa=2)
    add_box_text(label3, "勉強の仕方")

    study = doc.add_paragraph()
    sp(study, sb=2, sa=2)
    add_run(study, "君たちにやってほしいことは、扱った英文の完全な理解と音読による定着です。\n"
                   "それができているかを確認するテストにします。音読がんばって！！")

    # ======================================================
    # どこにあったの問題の例（Copenhagen版）
    # ======================================================
    ex_label = doc.add_paragraph()
    sp(ex_label, sb=10, sa=2)
    add_run(ex_label, "※どこにあったの問題の例", size=10)

    ex_inst = doc.add_paragraph()
    sp(ex_inst, sb=0, sa=2)
    add_run(ex_inst,
        "問　この英文中には下の６つの語句が抜けている。本来入るべき場所を指摘しなさい。\n"
        "　　解答欄には，それぞれの語句が入る場所の前後の単語を書きなさい"
        "（１つ目のみ，例として解答を示す）。", size=10)

    # 語句テーブル
    words = ["0. once", "1. so that", "2. such as", "3. quicker", "4. instead of", "5. well"]
    tbl = doc.add_table(rows=1, cols=len(words))
    tbl.style = 'Table Grid'
    for i, w in enumerate(words):
        cell = tbl.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(w)
        set_font(run, size=10)

    doc.add_paragraph()

    # 本文（語句削除済み）
    body_p1 = doc.add_paragraph()
    sp(body_p1, sb=0, sa=0, il=0.3)
    body_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(body_p1,
        "Surprisingly, Copenhagen's urban planners used to favor cars over other forms of "
        "transportation. The oil crisis of the 1970s led them to revise their transportation systems "
        "they would not have to rely on oil. Since then, the city has been installing cycling "
        "infrastructure the Green Wave and the Bicycle Snake. As a result, it is now to go downtown "
        "by bicycle than by car. That has led more people to cycle driving.", size=10)

    body_p2 = doc.add_paragraph()
    sp(body_p2, sb=0, sa=4, il=0.3)
    body_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(body_p2,
        "Making cities healthier and more attractive by adopting bike-friendly policies has come to "
        "be known as \"Copenhagenization.\" Even New York is following Copenhagen's example. It has "
        "started locating car parking spaces in the middle of the street. That approach is called the "
        "\"Protected Bicycle Lane\" policy and it prevents vehicles from entering or parking in the "
        "bike lane. Consequently, bicycle traffic has increased 1.6 times. Moreover, there are fewer "
        "accidents. Many other cities may \"Copenhagenize\" as they seek to build more sustainable, "
        "livable communities. The Copenhagen model is changing the world for the better.", size=10)

    # 解答欄
    ans_grid = doc.add_paragraph()
    sp(ans_grid, sb=0, sa=0)
    add_run(ans_grid,
        "1. （ planners ） once （ used ）　　　　2. （　　　　） so that （　　　　）\n"
        "3. （　　　　） such as （　　　　）　　4. （　　　　） quicker （　　　　）\n"
        "5. （　　　　） instead of （　　　　）　6. （　　　　） well （　　　　）", size=10)

    # ======================================================
    # 保存
    # ======================================================
    import os
    out = "/home/user/teacher-automation-lab/output/scope_comm2_final.docx"
    os.makedirs("/home/user/teacher-automation-lab/output", exist_ok=True)
    doc.save(out)
    print(f"保存完了: {out}")

if __name__ == "__main__":
    build_scope()
