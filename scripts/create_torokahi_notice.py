from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ページ設定（A4）
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# デフォルトフォントを游明朝に
doc.styles['Normal'].font.name = '游明朝'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

def para(text='', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=11, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = '游明朝'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    return p

# ── 日付・差出人（右寄せ）──────────────────────────────
para('令和８年　　月　　日', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_after=2)
para('駒澤大学高等学校　ラグビー部顧問', align=WD_ALIGN_PARAGRAPH.RIGHT, size=10.5)

# ── 宛先（左寄せ）──────────────────────────────────────
para('ラグビー部員　各位', size=11, space_before=4, space_after=12)

# ── 件名（中央・大きめ・太字）───────────────────────────
para('協会登録費の納入について（お願い）',
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=10)

# ── 前文 ───────────────────────────────────────────────
body1 = doc.add_paragraph()
body1.paragraph_format.space_after = Pt(8)
body1.paragraph_format.first_line_indent = Pt(11)
run1 = body1.add_run(
    '平素より、ラグビー部の活動にご理解・ご協力をいただき、誠にありがとうございます。'
    'さて、部活動を行うにあたり、東京都高等学校体育連盟ラグビーフットボール部への選手登録が必要となります。'
    'つきましては、下記のとおり協会登録費を徴収させていただきますので、'
    'ご確認のうえ、期日までにご提出くださいますようお願い申し上げます。'
)
run1.font.size = Pt(11)
run1.font.name = '游明朝'
run1._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

# ── 記 ────────────────────────────────────────────────
para('記', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_before=4, space_after=6)

# ── 詳細テーブル ────────────────────────────────────────
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

col_widths = [Cm(5.5), Cm(10.0)]
for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

def set_cell(row_idx, col_idx, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell = table.cell(row_idx, col_idx)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

set_cell(0, 0, '金　額', bold=True)
set_cell(0, 1, '２，０００円（登録費）')
set_cell(1, 0, '提出期限', bold=True)
set_cell(1, 1, '令和８年　　月　　日（　　）')
set_cell(2, 0, '提出先', bold=True)
set_cell(2, 1, '顧問（ホームルーム担任を通じて提出も可）')

# テーブル後の余白
para('', space_before=10, space_after=0)

# ── 注意書き ────────────────────────────────────────────
para('【お願い・提出方法】', bold=True, size=11, space_before=6, space_after=4)

notes = [
    '現金は必ず封筒に入れてご提出ください。',
    '封筒の表面に、①氏名　②クラス・出席番号　③金額（２，０００円）を明記してください。',
    '釣り銭のないようにご準備をお願いします。',
    '不明な点は、ラグビー部顧問までお問い合わせください。',
]
for note in notes:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(note)
    run.font.size = Pt(11)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

# ── 以上 ───────────────────────────────────────────────
para('以　上', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_before=12)

# ── 切り取り線 ─────────────────────────────────────────
para('', space_before=12)
cut_p = doc.add_paragraph()
cut_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cut_run = cut_p.add_run('✂  ─────────────────────────────────────  ✂')
cut_run.font.size = Pt(10)
cut_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── 封筒記載見本 ────────────────────────────────────────
para('【封筒の表面記載例】', bold=True, size=10.5, space_before=6, space_after=4)

env_table = doc.add_table(rows=1, cols=1)
env_table.style = 'Table Grid'
env_cell = env_table.cell(0, 0)
env_cell.width = Cm(8)

lines = [
    '氏名：　　　　　　　　',
    'クラス・番号：　　年　　組　　番',
    '金額：２，０００円',
    '',
    'ラグビー部　協会登録費',
]
env_p = env_cell.paragraphs[0]
env_p.paragraph_format.space_before = Pt(6)
env_p.paragraph_format.space_after  = Pt(6)
for line in lines:
    run = env_p.add_run(line + '\n')
    run.font.size = Pt(10.5)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

out_path = '/home/user/teacher-automation-lab/output/協会登録費徴収のお知らせ.docx'
import os; os.makedirs('/home/user/teacher-automation-lab/output', exist_ok=True)
doc.save(out_path)
print(f'保存完了: {out_path}')
