from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from camp_headcount_calc import compute

# ── 人数計算（共通モジュール camp_headcount_calc.py に一元化）───
rows, totals = compute()
total_b, total_l, total_d, total_s = totals['breakfast'], totals['lunch'], totals['dinner'], totals['stay']

# ── 文書生成 ──────────────────────────────────────────────
doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

doc.styles['Normal'].font.name = '游明朝'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

YELLOW = 'FFFF00'


def highlight_run(run, color=YELLOW):
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn('w:shd')):
        rPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    rPr.append(shd)


def para(text='', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=11,
         space_before=0, space_after=0, highlight=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = '游明朝'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
        if highlight:
            highlight_run(run)
    return p


# ── FAX送信状 ──────────────────────────────────────────
para('FAX送信状', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=10)

fax_table = doc.add_table(rows=6, cols=2)
fax_table.style = 'Table Grid'
fax_col_widths = [Cm(4.5), Cm(10.9)]
for row in fax_table.rows:
    for j, cell in enumerate(row.cells):
        cell.width = fax_col_widths[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)


def set_fax_cell(row_idx, col_idx, text, bold=False):
    cell = fax_table.cell(row_idx, col_idx)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')


set_fax_cell(0, 0, '送信日', bold=True)
set_fax_cell(0, 1, '2026年7月17日')
set_fax_cell(1, 0, '宛先', bold=True)
set_fax_cell(1, 1, '有限会社　菅平ホテル　御中（FAX：0268-74-3548）')
set_fax_cell(2, 0, '発信元', bold=True)
set_fax_cell(2, 1, '駒澤大学高等学校　ラグビー部顧問　占部涼也')
set_fax_cell(3, 0, '発信元連絡先', bold=True)
set_fax_cell(3, 1, 'Tel（携帯）：080-5032-9150 / Mail：ryoyaurabe@gmail.com')
set_fax_cell(4, 0, '件名', bold=True)
set_fax_cell(4, 1, '夏季合宿　参加人数確定のご連絡')
set_fax_cell(5, 0, '枚数', bold=True)
set_fax_cell(5, 1, '本紙を含め　　枚', bold=False)
highlight_run(fax_table.cell(5, 1).paragraphs[0].runs[0])

doc.add_page_break()

# ── 日付・差出人（右寄せ）──────────────────────────────
para('2026年7月17日', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_after=2)
para('駒澤大学高等学校　ラグビー部顧問　占部涼也', align=WD_ALIGN_PARAGRAPH.RIGHT, size=10.5)

# ── 宛先（左寄せ）──────────────────────────────────────
para('有限会社　菅平ホテル　御中', size=11, space_before=4, space_after=12)

# ── 件名 ───────────────────────────────────────────────
para('夏季合宿　参加人数確定のご連絡', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=10)

# ── 前文 ───────────────────────────────────────────────
body1 = doc.add_paragraph()
body1.paragraph_format.space_after = Pt(8)
body1.paragraph_format.first_line_indent = Pt(11)
run1 = body1.add_run(
    '拝啓　時下ますますご清栄のこととお慶び申し上げます。平素より大変お世話になっております。'
    '2026年6月9日付でいただきましたご請求書（宿泊人数集計表）につきまして、'
    '参加生徒の内訳が確定いたしましたので、下記のとおりご連絡申し上げます。'
    'お手数をおかけいたしますが、お食事のご準備の調整にお役立ていただけますと幸いです。'
)
run1.font.size = Pt(11)
run1.font.name = '游明朝'
run1._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

# ── 記 ────────────────────────────────────────────────
para('記', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_before=4, space_after=6)

para('１．合宿期間', bold=True, size=11, space_before=4, space_after=2)
para('　　8月1日（土）〜8月7日（金）　6泊7日', size=11, space_after=8)

# ── 内訳テーブル ────────────────────────────────────────
para('２．参加人数の内訳', bold=True, size=11, space_before=4, space_after=2)

table1 = doc.add_table(rows=5, cols=2)
table1.style = 'Table Grid'
col_widths = [Cm(6.0), Cm(9.4)]
for row in table1.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)


def set_cell(table, row_idx, col_idx, text, bold=False, highlight=False):
    cell = table.cell(row_idx, col_idx)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    if highlight:
        highlight_run(run)


set_cell(table1, 0, 0, '部員（生徒）', bold=True)
set_cell(table1, 0, 1, '42名　（選手39名・マネージャー3名）')
set_cell(table1, 1, 0, '引率顧問', bold=True)
set_cell(table1, 1, 1, '2名　（顧問1：全日程帯同／顧問2：8月5日夕食〜合流）')
set_cell(table1, 2, 0, '外部コーチ', bold=True)
set_cell(table1, 2, 1, '2名　（入れ替わりで常時1名帯同。1人目：8/1〜8/2、2人目：8/3〜8/5）')
set_cell(table1, 3, 0, '', bold=True)
set_cell(table1, 3, 1, '　8/6・8/7はコーチの帯同なし')
set_cell(table1, 4, 0, '最大同時人数', bold=True)
set_cell(table1, 4, 1, '45名　（部員42名＋顧問2名＋コーチ1名。8/5夕食〜宿泊）')

para('', space_before=8)

para('３．途中合流する部員・顧問について', bold=True, size=11, space_before=4, space_after=2)
mid_join = [
    '椎名　薫（2年F組・マネージャー）：8月4日午後〜合流（夕食からカウント）',
    '吉田　青空（1年E組・選手）：8月5日午後〜合流（夕食からカウント）',
    '顧問2　占部涼也：8月5日午後〜合流（夕食からカウント）',
]
for line in mid_join:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

para('', space_before=8)

para('４．外部コーチの帯同について', bold=True, size=11, space_before=4, space_after=2)
coach_sched = [
    '1人目：8月1日〜8月2日（8/2の夕食・宿泊まで帯同し、翌朝までに交代）',
    '2人目：8月3日〜8月5日（8/5の夕食・宿泊まで帯同）',
    '8月6日・8月7日はコーチの帯同なし',
]
for line in coach_sched:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

para('', space_before=8)

# ── 日別人数表 ────────────────────────────────────────
para('５．日別人数表（部員・引率顧問・外部コーチの合計）', bold=True, size=11, space_before=4, space_after=2)

table2 = doc.add_table(rows=len(rows) + 2, cols=5)
table2.style = 'Table Grid'
col_widths2 = [Cm(3.6), Cm(2.9), Cm(2.9), Cm(2.9), Cm(2.9)]
for row in table2.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_widths2[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

headers = ['日付', '朝食', '昼食', '夕食', '宿泊']
for j, h in enumerate(headers):
    set_cell(table2, 0, j, h, bold=True)

for i, (day, b, l, d, s) in enumerate(rows):
    set_cell(table2, i + 1, 0, day, bold=True)
    set_cell(table2, i + 1, 1, str(b))
    set_cell(table2, i + 1, 2, str(l))
    set_cell(table2, i + 1, 3, str(d))
    set_cell(table2, i + 1, 4, str(s))

last = len(rows) + 1
set_cell(table2, last, 0, '延べ人数', bold=True)
set_cell(table2, last, 1, str(total_b), bold=True)
set_cell(table2, last, 2, str(total_l), bold=True)
set_cell(table2, last, 3, str(total_d), bold=True)
set_cell(table2, last, 4, str(total_s), bold=True)

para('', space_before=8)

note2 = doc.add_paragraph()
note2.paragraph_format.space_after = Pt(8)
run_n2 = note2.add_run(
    '※　上表は部員42名（途中合流2名を含む）・引率顧問2名（顧問2は8/5夕食〜）・'
    '外部コーチ（8/1〜8/5、入れ替わりで常時1名）を合算した確定人数です。'
    'ご請求書（6/9付）でいただいた46名／日の一律見積と比べ、日によって人数が変動いたしますので、'
    '上表の人数でご準備をお願いいたします。'
)
run_n2.font.size = Pt(9.5)
run_n2.font.name = '游明朝'
run_n2._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

para('６．お願い', bold=True, size=11, space_before=4, space_after=2)
body3 = doc.add_paragraph()
body3.paragraph_format.space_after = Pt(8)
run3 = body3.add_run(
    '上記人数にもとづき、お食事の提供数調整をお願いできますと幸いです。'
    'なお、最終のご請求額につきましては、合宿終了後の実績人数にもとづき、'
    '精算いただけますようお願い申し上げます。ご不明な点がございましたら、下記までご連絡ください。'
)
run3.font.size = Pt(11)
run3.font.name = '游明朝'
run3._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

para('敬具', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_before=8, space_after=4)

# ── 連絡先 ────────────────────────────────────────────
para('【本状に関するお問い合わせ先】', bold=True, size=10.5, space_before=8, space_after=4)
contact_lines = [
    '駒澤大学高等学校　ラグビー部顧問　占部涼也',
    'Mail：ryoyaurabe@gmail.com',
    'Tel（携帯）：080-5032-9150',
]
for line in contact_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(line)
    run.font.size = Pt(10.5)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

para('以　上', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_before=12)

out_path = '/home/user/teacher-automation-lab/output/2026_夏合宿_人数確定のご連絡_菅平ホテル様.docx'
os.makedirs('/home/user/teacher-automation-lab/output', exist_ok=True)
doc.save(out_path)
print(f'保存完了: {out_path}')
print(f'延べ朝食={total_b} 延べ昼食={total_l} 延べ夕食={total_d} 延べ宿泊={total_s}')
