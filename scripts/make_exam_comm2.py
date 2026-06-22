"""
英語コミュニケーションII 1学期期末試験 Word生成スクリプト
過去問（1学期中間試験）のレイアウト・書式を完全に模倣して .docx を出力する
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== ヘルパー関数 ==========

def set_font(run, name_ja="MS Mincho", name_en="Times New Roman", size=10.5, bold=False, italic=False, underline=False):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_ja)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def highlight_yellow(run):
    """語句・短答の模範解答用：黄色ハイライト"""
    r = run._r
    rPr = r.get_or_add_rPr()
    existing = rPr.find(qn('w:highlight'))
    if existing is not None:
        rPr.remove(existing)
    hl = OxmlElement('w:highlight')
    hl.set(qn('w:val'), 'yellow')
    rPr.append(hl)

def add_char_border(run):
    """長文読解の選択肢解答用：文字を枠線で囲む"""
    r = run._r
    rPr = r.get_or_add_rPr()
    bdr = OxmlElement('w:bdr')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '4')
    bdr.set(qn('w:space'), '1')
    bdr.set(qn('w:color'), 'auto')
    rPr.append(bdr)

def add_run(para, text, bold=False, italic=False, underline=False, size=10.5,
            ja_font="MS Mincho", en_font="Times New Roman", yellow=False, boxed=False):
    run = para.add_run(text)
    set_font(run, name_ja=ja_font, name_en=en_font, size=size, bold=bold, italic=italic, underline=underline)
    if yellow:
        highlight_yellow(run)
    if boxed:
        add_char_border(run)
    return run

def set_para_format(para, align=None, space_before=0, space_after=0,
                    line_spacing=None, indent_left=0, indent_first=0,
                    right_tab_at=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.left_indent  = Cm(indent_left)
    pf.first_line_indent = Cm(indent_first)
    if align is not None:
        para.alignment = align
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)
    if right_tab_at is not None:
        pf.tab_stops.add_tab_stop(Cm(right_tab_at), WD_TAB_ALIGNMENT.RIGHT)

# 本文幅（A4・左右2.5cm余白）に対する右端タブ位置
RIGHT_MARGIN_TAB = 16.0

def add_boxed_number(para, number):
    """実際の枠線で囲んだ半角数字（過去問の□1スタイル）"""
    run = para.add_run(str(number))
    set_font(run, size=11, bold=True)
    r = run._r
    rPr = r.get_or_add_rPr()
    bdr = OxmlElement('w:bdr')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '8')
    bdr.set(qn('w:space'), '2')
    bdr.set(qn('w:color'), 'auto')
    rPr.append(bdr)
    return run

def add_section_header(doc, number, text, points=None):
    """大問ヘッダー: [1] 指示文 （XX点）"""
    para = doc.add_paragraph()
    set_para_format(para, space_before=8, space_after=2)
    add_boxed_number(para, number)
    add_run(para, "　" + text, bold=True, size=10.5)
    if points:
        add_run(para, f"（{points}点）", bold=True, size=10.5)
    return para

def add_dialogue_line(doc, speaker, body_runs):
    """ぶら下げインデントの会話文（枠線なし）
    body_runs: [(text, kwargs), ...]
    """
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=0,
                    indent_left=1.0, indent_first=-1.0)
    add_run(p, speaker)
    for text, kwargs in body_runs:
        add_run(p, text, **kwargs)
    return p

def set_page_layout(doc):
    """A4・余白設定"""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ========== SVOCM パーサー・レンダラー ==========

def parse_svocm(text):
    """SVOCM記法の文字列を (content, tag) のリストに変換する。
    tag: 'S'=□囲み, 'V'=下線, 'O'=二重下線, 'C'=波線, None=平文
    M(...)・M（...）は"M"プレフィックスを除去して平文として返す。
    """
    result = []
    i = 0
    n = len(text)
    while i < n:
        found = False
        for tag in ['S', 'V', 'O', 'C']:
            for ob, cb in [('[', ']'), ('〈', '〉')]:
                if text[i:i+2] == f'{tag}{ob}':
                    depth = 1
                    j = i + 2
                    while j < n and depth > 0:
                        if text[j] == ob:
                            depth += 1
                        elif text[j] == cb:
                            depth -= 1
                        j += 1
                    result.append((text[i+2:j-1], tag))
                    i = j
                    found = True
                    break
            if found:
                break
        if not found:
            # M(...)・M（...）: Mを除いてカッコごと平文
            m_found = False
            for ob, cb in [('(', ')'), ('（', '）')]:
                if text[i:i+2] == f'M{ob}':
                    depth = 1
                    j = i + 1
                    while j < n and depth > 0:
                        if text[j] == ob:
                            depth += 1
                        elif text[j] == cb:
                            depth -= 1
                        j += 1
                    result.append((text[i+1:j], None))
                    i = j
                    m_found = True
                    break
            if not m_found:
                # 平文：次のタグ開始まで
                j = i + 1
                while j < n:
                    hit = False
                    for tag in ['S', 'V', 'O', 'C']:
                        for ob in ['[', '〈']:
                            if text[j:j+2] == f'{tag}{ob}':
                                hit = True
                                break
                        if hit:
                            break
                    for ob in ['(', '（']:
                        if text[j:j+2] == f'M{ob}':
                            hit = True
                            break
                    if hit:
                        break
                    j += 1
                result.append((text[i:j], None))
                i = j
    return result


def add_svocm_text(para, analysis_str, size=9.5, yellow=False):
    """SVOCM記法文字列を視覚書式付きでパラグラフに追加する。"""
    for content, tag in parse_svocm(analysis_str):
        run = para.add_run(content)
        set_font(run, size=size)
        if yellow:
            highlight_yellow(run)
        r = run._r
        rPr = r.get_or_add_rPr()
        if tag == 'S':
            bdr = OxmlElement('w:bdr')
            bdr.set(qn('w:val'), 'single')
            bdr.set(qn('w:sz'), '4')
            bdr.set(qn('w:space'), '1')
            bdr.set(qn('w:color'), 'auto')
            rPr.append(bdr)
        elif tag == 'V':
            run.font.underline = True
        elif tag == 'O':
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'double')
            rPr.append(u)
        elif tag == 'C':
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'wave')
            rPr.append(u)


# ========== 本体 ==========

def build_exam(student=False):
    doc = Document()
    set_page_layout(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

    # =========================================================
    # タイトル
    # =========================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(title, space_before=0, space_after=4)
    add_run(title, "２０２６年　２学年受験コース　英語コミュニケーションⅡ　１学期期末試験",
            bold=True, size=12)

    doc.add_paragraph()

    # =========================================================
    # 大問１：空所補充（過去問の出典文と重複しない独立した9文）
    # =========================================================
    add_section_header(doc, 1,
        "日本文の意味になるように英文の空所に適語を入れなさい。"
        "空欄にアルファベットがあるものは、それではじまる単語を答えること。",
        points="９")

    q1_data = [
        # (1) L3 Part3①(6)
        ("この地域の多くの人は名古屋に", "通勤している", "。",
         "Many people in this area （　　　　　　　　　） to Nagoya.",
         "commute"),
        # (2) L4 Part4①(6)
        ("平和への願いは", "万国共通", "です。",
         "Hope for peace is （　　　　　　　　）.",
         "universal"),
        # (3) L3 Grammar G1③(2)
        ("鍵が見つからない。家の中に", "置いてきてしまったにちがいない", "。",
         "I can't find my key. I （　　　）（　　　　　）（　　　） it in the house.",
         "must / have / left"),
        # (4) L3 Part1①(5)
        ("私のおじは農家で、", "温室", "でイチゴを育てている。",
         "My uncle is a farmer and he grows strawberries in a （　　　　　　　　　）.",
         "greenhouse"),
        # (5) L4 Part1①(8)
        ("私にはこの英語の語の", "ニュアンス", "がわからない。",
         "I don't understand the （　　　　　　　　　） of this English word.",
         "nuance"),
        # (6) L4 Part4①(2)
        ("生徒会は全生徒に", "調査を行う", "つもりです。",
         "The student council is going to （　　　　　　　） a survey of all the students.",
         "conduct"),
        # (7) L4 Grammar G2③(1)
        ("睡眠不足が健康に悪いことは", "言うまでもない", "。",
         "It （　　　　）（　　　　　　　） saying that lack of sleep is bad for your health.",
         "goes / without"),
        # (8) L4 Part1①(1) ―「顔認識」のうち recognition を与え、facial を答えさせる
        ("そのシステムは", "顔", "認識技術を使用している。",
         "The system uses （　　　　　　） recognition technology.",
         "facial"),
        # (9) L3 Part3①(4)
        ("彼らはキャンパーたちが山から", "安全に", "戻ってくるのを助けた。",
         "They helped the campers come back from the mountain （　　　　　　）.",
         "safely"),
    ]

    for i, (pre, underline_text, post, english_line, hint) in enumerate(q1_data, 1):
        jp_para = doc.add_paragraph()
        set_para_format(jp_para, space_before=3, space_after=0, indent_left=0.3,
                        right_tab_at=RIGHT_MARGIN_TAB)
        add_run(jp_para, f"({i})　{pre}")
        add_run(jp_para, underline_text, underline=True)
        add_run(jp_para, post)
        if not student:
            add_run(jp_para, "\t")
            add_run(jp_para, hint, bold=True, yellow=True)

        en_para = doc.add_paragraph()
        set_para_format(en_para, align=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=0, space_after=4, indent_left=0.8)
        add_run(en_para, english_line)

    doc.add_paragraph()

    # =========================================================
    # 大問２：読解＋設問（枠線なし・ぶら下げインデント）
    # =========================================================
    add_section_header(doc, 2,
        "Read the following text and answer the questions.", points="６")

    add_dialogue_line(doc, "Student: ",
        [("I heard that Copenhagen has become the world's most bicycle-friendly city.", {})])
    add_dialogue_line(doc, "Mr. Hansen: ", [
        ("That's right. ", {}),
        ("①", {}),
        ("The city has implemented some unique policies with the goal of becoming the best bike city. ", {}),
        ("②", {}),
        ("Let me explain one of them—the Green Wave. As long as cyclists maintain a speed of 20 kilometers per hour, the traffic lights are green all the way during rush hour on major bicycle routes. ", {}),
        ("③", {}),
        ("Copenhagen is also planning to expand its international airport to attract more business travelers. ", {}),
        ("④", {}),
        ("The road signals are operated simultaneously with the movements of the cyclists, which helped eliminate traffic congestion.", {}),
    ])
    add_dialogue_line(doc, "Student: ",
        [("Does everyone have to cycle at exactly 20 km/h?", {})])
    add_dialogue_line(doc, "Mr. Hansen: ", [
        ("Well, ( A ) some cyclists would rather go faster, but they can be a danger to other cyclists during rush hour. "
         "The Green Wave makes them travel at a safe speed, because if they go too fast, they'll get caught at a red light. "
         "( B ), everyone cycles at the safest speed without sacrificing efficiency.", {}),
    ])

    doc.add_paragraph()

    q2_p1 = doc.add_paragraph()
    set_para_format(q2_p1, space_before=2, space_after=2, indent_left=0.3,
                    right_tab_at=RIGHT_MARGIN_TAB)
    add_run(q2_p1, "1. Choose one sentence from ① to ④ ", bold=True)
    add_run(q2_p1, "that does not belong in the paragraph.", bold=True, underline=True)
    if not student:
        add_run(q2_p1, "\t")
        add_run(q2_p1, "③", bold=True, yellow=True)

    q2_p2 = doc.add_paragraph()
    set_para_format(q2_p2, space_before=2, space_after=2, indent_left=0.3)
    add_run(q2_p2, "2. Choose the best option for the blanks A and B.", bold=True)
    add_run(q2_p2, "（Words in the options ")
    add_run(q2_p2, "begin with lowercase letters,", underline=True)
    add_run(q2_p2, " even when they come at the beginning of a sentence.）")

    q2_p3 = doc.add_paragraph()
    set_para_format(q2_p3, space_before=0, space_after=2, indent_left=0.5,
                    right_tab_at=RIGHT_MARGIN_TAB)
    add_run(q2_p3, "[ ア it's true that ／ イ in this way ／ ウ similarly ／ エ what's more ／ オ nevertheless ]")
    if not student:
        add_run(q2_p3, "\t")
        add_run(q2_p3, "A", bold=True)
        add_run(q2_p3, "ア", bold=True, yellow=True)
        add_run(q2_p3, " B", bold=True)
        add_run(q2_p3, "イ", bold=True, yellow=True)

    # =========================================================
    # 大問３：語句選択（ア〜エ）Lesson 4（絵文字）Part 4 使用
    # =========================================================
    add_section_header(doc, 3,
        "Choose the best word for each blank from options ア to エ. 記号で答えなさい。",
        points="４")

    q3_opt = doc.add_paragraph()
    set_para_format(q3_opt, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
    add_run(q3_opt, "ア recognized　　　イ baffled　　　ウ conducted　　　エ rated")

    q3_body = doc.add_paragraph()
    set_para_format(q3_body, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=2, space_after=2)
    add_run(q3_body,
        "Lastly, people need to have seen some emoticons before they can understand their meanings. "
        "An experiment was （　1　） in Japan, Cameroon, and Tanzania. "
        "The participants （　2　） three different styles of emoticons using a "
        "'sadness-happiness' scale. "
        "The Japanese participants easily （　3　） the emotions in the emoticons. "
        "On the other hand, the participants in Cameroon and Tanzania hardly understood them. "
        "They were utterly （　4　） by what the shape of each emoticon indicated. "
        "In other words, emoticons do not look like facial expressions to everyone.")

    q3_ans = doc.add_paragraph()
    set_para_format(q3_ans, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=6)
    if not student:
        add_run(q3_ans, "１ウ　２エ　３ア　４イ", bold=True, yellow=True)

    # =========================================================
    # 大問４：語句挿入（Lesson 3 Part 4 Copenhagenization 使用）
    # =========================================================
    add_section_header(doc, 4,
        "この英文中には下の５つの語句が抜けている。本来入るべき場所を指摘しなさい。"
        "解答欄には，それぞれの語句が入る場所の前後の単語を書きなさい。"
        "No.０は解答例である。１〜５を解答しなさい。",
        points="１０")

    word_table = doc.add_table(rows=1, cols=6)
    word_table.style = 'Table Grid'
    words = ["0. once", "1. so that", "2. such as", "3. quicker", "4. from", "5. well"]
    for i, word in enumerate(words):
        cell = word_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, word, size=10)

    doc.add_paragraph()

    ex_para = doc.add_paragraph()
    set_para_format(ex_para, space_before=0, space_after=2, indent_left=0.3)
    add_run(ex_para, "解答例　　（ planners ） once （ used ）")

    # 語句挿入本文（Lesson 3 Part 4 そのまま・0〜5の語を削除）
    body_para1 = doc.add_paragraph()
    set_para_format(body_para1, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=2, space_after=0)
    add_run(body_para1,
        "Surprisingly, Copenhagen's urban planners used to favor cars over other forms of "
        "transportation. The oil crisis of the 1970s led them to revise their transportation systems "
        "they would not have to rely on oil. Since then, the city has been installing cycling "
        "infrastructure the Green Wave and the Bicycle Snake. As a result, it is now to go downtown "
        "by bicycle than by car. That has led more people to cycle instead of driving.")

    body_para2 = doc.add_paragraph()
    set_para_format(body_para2, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=4)
    add_run(body_para2,
        "Making cities healthier and more attractive by adopting bike-friendly policies has come to "
        "be known as \"Copenhagenization.\" Even New York is following Copenhagen's example. It has "
        "started locating car parking spaces in the middle of the street. That approach is called the "
        "\"Protected Bicycle Lane\" policy and it prevents vehicles entering or parking in the "
        "bike lane. Consequently, bicycle traffic has increased 1.6 times. Moreover, there are fewer "
        "accidents. Many other cities may \"Copenhagenize\" as they seek to build more sustainable, "
        "livable communities. The Copenhagen model is changing the world for the better.")

    ans4_para = doc.add_paragraph()
    set_para_format(ans4_para, space_before=2, space_after=6, indent_left=0.3)
    ans4_items = [
        "1.（systems）so that（they）",
        "2.（infrastructure）such as（the）",
        "3.（now）quicker（to）",
        "4.（vehicles）from（entering）",
        "5.（may）well（\"Copenhagenize\"）",
    ]
    if not student:
        add_run(ans4_para, "　　".join(ans4_items), bold=True, yellow=True, size=9.5)

    # =========================================================
    # 大問５：英文解釈（入門英文問題精講 33,37,38,41 より）
    # =========================================================
    add_section_header(doc, 5, "以下の英文解釈に関する問題に答えなさい。（各4点：構造2点・和訳2点）", points="１６")

    # Part 1 指示文
    inst1 = doc.add_paragraph()
    set_para_format(inst1, space_before=4, space_after=2, indent_left=0.3)
    add_run(inst1,
        "下の例にならい、次の文１〜２に SVOCM の記号を振り構造を表しなさい。また、それにもとづいた英文の"
        "日本語訳を書きなさい。ただし、1 つの単語につき 2 つ以上の記号・下線は付さなくて良い。"
        "文全体の大きな構造がわかるように図示すること。")

    ex_label = doc.add_paragraph()
    set_para_format(ex_label, space_before=2, space_after=0, indent_left=0.3)
    add_run(ex_label, "例", bold=True, size=9)

    ex_key = doc.add_paragraph()
    set_para_format(ex_key, space_before=0, space_after=2, indent_left=0.5)
    add_run(ex_key,
        "（記号）S → 四角　、V → 下線　、O → 二重下線　、C → 波線　、M → （まるカッコ）", size=9)

    ex_sent = doc.add_paragraph()
    set_para_format(ex_sent, space_before=0, space_after=6, indent_left=0.5)
    add_svocm_text(ex_sent,
        "S[Developing nations] V[were] C[under-represented] "
        "M(, with competitors facing obstacles ranging from prejudice against disability "
        "to the high cost of wheelchairs).", size=9.5)

    def add_char_border_run(para, text, size=9.5):
        run = para.add_run(text)
        set_font(run, size=size)
        highlight_yellow(run)
        return run

    problems_p1 = [
        (
            "Becoming an adult is a step-by-step process, and just when the young are finally "
            "wise enough to be treated as young adults is not the time to give them free access "
            "to the drinks bar.",
            "S〈Becoming an adult〉 V[is] C[a step-by-step process], and "
            "S〈just when the young are finally wise enough to be treated as young adults〉 "
            "V[is] not C[the time [to give them free access to the drinks bar]].",
            "大人になることは段階的なプロセスであり、若者たちがようやく若い大人として扱われるほど賢くなった"
            "まさにその時こそ、彼らにお酒を自由に飲ませる時ではない。",
            None, None
        ),
        (
            "Although we know that the earth revolves around the sun, we cannot recite "
            "the astronomical observations and calculations that led to that conclusion.",
            "M(Although S[we] V[know] O〈that the earth revolves around the sun〉), "
            "S[we] V[cannot recite] O[the astronomical observations and calculations "
            "[that led to that conclusion]].",
            "地球が太陽の周りを公転していることは知っていても、私たちはその結論に至った天文学的な観測や"
            "計算を暗唱することはできない。",
            None, None
        ),
    ]

    for i, (sentence, answer, translation, sub_q, sub_ans) in enumerate(problems_p1, 1):
        p_num = doc.add_paragraph()
        set_para_format(p_num, space_before=6, space_after=0, indent_left=0.3)
        add_run(p_num, f"　{i}．", bold=True)
        if isinstance(sentence, list):
            for text, ul in sentence:
                add_run(p_num, text, underline=ul)
        else:
            add_run(p_num, sentence)

        if not student:
            label_s = doc.add_paragraph()
            set_para_format(label_s, space_before=1, space_after=0, indent_left=0.5)
            add_run(label_s, "（記号）", size=9.5)
            blank_s = doc.add_paragraph()
            set_para_format(blank_s, space_before=0, space_after=0, indent_left=0.5)
            add_svocm_text(blank_s, answer, size=9.5, yellow=True)
            label_t = doc.add_paragraph()
            set_para_format(label_t, space_before=2, space_after=0, indent_left=0.5)
            add_run(label_t, "（訳）", size=9.5)
            blank_t = doc.add_paragraph()
            set_para_format(blank_t, space_before=0, space_after=0, indent_left=0.5)
            add_run(blank_t, translation, yellow=True, size=9.5)

        doc.add_paragraph()

    # Part 2 指示文（正誤判定）
    inst2 = doc.add_paragraph()
    set_para_format(inst2, space_before=8, space_after=2, indent_left=0.3)
    add_run(inst2,
        "下の３〜６の構文解釈が正しければア、間違っていればイを書きなさい。"
        "３〜４は SVOCM について、５〜６は句・節について（記号は"
        "〈名詞句・節〉　［形容詞句・節］　（副詞句・節））である。")

    problems_p2 = [
        ("3.", "S[X-rays] V[allowed] O[them] C[to look into their patients, "
               "identify where there were problems, and cure them].", "ア",
         "ア（正しい）：allow O to do = SVOC の構造。"),
        ("4.", "S[Willpower] V[is] C[a mysterious] M（force [that helps us control "
               "our actions and achieve our goals]）.", "イ",
         "イ（間違い）：C は「a mysterious force [that...]」全体。"),
        ("5.", "S[Chopik] V[says] O〈he isn't suggesting 〈we ignore our families〉〉, "
               "but O〈that friends make us feel better〉.", "ア",
         "ア（正しい）：says に対して二つの目的語節が but で並列されている。"),
        ("6.", "M（With friends）S[you] V[are] C[more likely to do activities] "
               "M（— they provide an outlet）.", "イ",
         "イ（間違い）：ダッシュ以降の「they provide an outlet」は独立した節。"),
    ]

    sentences_p2 = [
        "X-rays allowed them to look into their patients, identify where there were problems, and cure them.",
        "Willpower is a mysterious force that helps us control our actions and achieve our goals.",
        "Chopik says he isn't suggesting we ignore our families, but that friends make us feel better.",
        "With friends you are more likely to do activities — they provide an outlet.",
    ]

    for (num, analysis, ans, explanation), sentence in zip(problems_p2, sentences_p2):
        p_sent = doc.add_paragraph()
        set_para_format(p_sent, space_before=4, space_after=0, indent_left=0.3,
                        right_tab_at=RIGHT_MARGIN_TAB)
        add_run(p_sent, f"　{num}　", bold=True)
        add_run(p_sent, sentence)
        if not student:
            add_run(p_sent, "\t")
            add_run(p_sent, ans, bold=True, yellow=True)

        p_ana = doc.add_paragraph()
        set_para_format(p_ana, space_before=1, space_after=0, indent_left=0.7)
        add_svocm_text(p_ana, analysis, size=9.5)

        if not student:
            p_ans = doc.add_paragraph()
            set_para_format(p_ans, space_before=1, space_after=4, indent_left=0.7)
            add_run(p_ans, "　" + explanation, size=9, yellow=True)

    if not student:
        ans_row = doc.add_paragraph()
        set_para_format(ans_row, space_before=6, space_after=6,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(ans_row, "３ア　４イ　５ア　６イ", bold=True, yellow=True)

    # =========================================================
    # 大問６：並べ替え（動画でわかる英文法 例文51〜60より・10点）
    # =========================================================
    add_section_header(doc, 6,
        "日本語の意味を表す英文になるように（　）内の語句を並べ替え，（　）内で２番目と５番目"
        "に来る語句を解答欄に記入しなさい。解答欄には１語句ずつ記入すること。"
        "なお，複数の単語からなる選択肢も一つの語句として数える。"
        "また，文頭に来る語も１文字目は小文字になっている。",
        points="１０")

    q6_items = [
        (
            "定期的に運動する人々は自分自身により満足している。",
            "( those / who / exercise regularly / feel / better / about / themselves ).",
            "who", "better",
            "Those who exercise regularly feel better about themselves."
        ),
        (
            "今日できることを明日に延期するな。",
            "Don't ( put / off / until / tomorrow / what / you / can / do ) today.",
            "off", "what",
            "Don't put off until tomorrow what you can do today."
        ),
        (
            "彼女は、その賞を取ると私が思っている女の子だ。",
            "She ( is / the / girl / who / I / think / will / win ) the prize.",
            "the", "I",
            "She is the girl who I think will win the prize."
        ),
        (
            "十分な睡眠をとることが、あなたに必要なことだ。",
            "( what / you / need / is to get / enough / sleep ).",
            "you", "enough",
            "What you need is to get enough sleep."
        ),
        (
            "新しい経験を共有することができる人たちと友だちになることが重要だ。",
            "It is important ( to / make / friends / with / people / with / whom / you ) can share new experiences.",
            "make", "people",
            "It is important to make friends with people with whom you can share new experiences."
        ),
    ]

    for i, (jp, en, ans4th, ans8th, full_sentence) in enumerate(q6_items, 1):
        jp_p = doc.add_paragraph()
        set_para_format(jp_p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=4, space_after=0,
                        indent_left=0.3, right_tab_at=RIGHT_MARGIN_TAB)
        add_run(jp_p, f"　{i}．{jp}")
        if not student:
            add_run(jp_p, "\t")
            add_run(jp_p, f"2番目:{ans4th} / 5番目:{ans8th}", bold=True, yellow=True)

        en_p = doc.add_paragraph()
        set_para_format(en_p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
                        space_after=4 if student else 0, indent_left=0.8)
        add_run(en_p, en)

        if not student:
            ans_p = doc.add_paragraph()
            set_para_format(ans_p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, indent_left=0.8)
            add_run(ans_p, full_sentence, size=9.5)

    # =========================================================
    # 大問７：初見長文読解「シェイクスピアの言語」（英検2級レベル・15点）
    # =========================================================
    add_section_header(doc, 7,
        "以下の英文を読み、各問いに答えなさい。番号で答えなさい。",
        points="１５")

    ltitle = doc.add_paragraph()
    ltitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(ltitle, space_before=2, space_after=2)
    add_run(ltitle, "Shakespeare and the English Language", bold=True)

    reading_paras = [
        "　William Shakespeare, who lived from 1564 to 1616, is widely considered the greatest writer in the English language. His plays and poems have been translated into every major language and are performed more often than those of any other playwright. However, many people do not realize how deeply Shakespeare has influenced the way we speak English today.",
        "　During Shakespeare's time, the English language was going through a period of rapid change known as Early Modern English. New words were being created at an extraordinary rate to describe new ideas, places, and inventions. Shakespeare himself is believed to have invented more than 1,700 words that are still used in modern English. Words such as \"bedroom,\" \"lonely,\" \"generous,\" and \"fashionable\" all appear for the first time in Shakespeare's works.",
        "　Shakespeare also contributed many phrases that remain common expressions today. When someone says \"break the ice\" to describe starting a conversation with a stranger, or \"all that glitters is not gold\" as a warning not to judge by appearances, they are quoting Shakespeare — often without knowing it.",
        "　Readers today sometimes struggle with Shakespeare's language because it contains archaic forms. For example, Shakespeare often used \"thou\" and \"thee\" instead of \"you,\" and verb endings like \"-eth\" and \"-est\" are no longer part of standard English. Despite these differences, studying Shakespeare's original language can give us a unique window into how English has changed over the past 400 years.",
        "　Today, Shakespeare's influence can be felt not only in literature but also in everyday conversation. Understanding his language helps us appreciate the rich history of English and the remarkable creativity of one of its greatest contributors.",
    ]

    for rp in reading_paras:
        p = doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0, line_spacing=18)
        add_run(p, rp)

    doc.add_paragraph()

    q7_items = [
        ("(1)", "What is one thing that surprises many people about Shakespeare?", "3",
         ["He lived during a period when printed books were not yet widely available to ordinary people.",
          "His plays are now staged far less frequently than those of other well-known dramatists worldwide.",
          "He had a deep influence on the words and expressions we use in English today.",
          "He personally translated many of his own plays and poems into several foreign languages."]),
        ("(2)", "Which of the following is NOT mentioned as a word invented by Shakespeare?", "4",
         ["bedroom",
          "fashionable",
          "generous",
          "extraordinary"]),
        ("(3)", "What does \"break the ice\" mean, according to the passage?", "2",
         ["To break through a layer of ice that has formed on the surface of a frozen lake or river.",
          "To begin a conversation with a person whom one has not met or spoken to before.",
          "To perform a memorable scene from one of Shakespeare's famous plays in front of an audience.",
          "To caution someone who is about to judge a situation based only on its outward appearance."]),
        ("(4)", "Why do modern readers sometimes find Shakespeare's language difficult?", "3",
         ["His plays include many invented words and expressions not recorded in any modern dictionary.",
          "The storylines and plots of his plays are far too complex for today's general readers to follow.",
          "He often used old-fashioned pronouns and verb endings that are no longer part of standard English.",
          "His original manuscripts were composed in a script that modern English speakers cannot read at all."]),
        ("(5)", "What is the main idea of this passage?", "3",
         ["Shakespeare's plays and poems are performed and translated more widely than any other writer's works.",
          "Early Modern English, spoken in Shakespeare's era, differed greatly from the language we use today.",
          "Shakespeare made lasting contributions to the English language that continue to influence us.",
          "Reading classic literature such as Shakespeare is considered the most reliable path to improving English."]),
    ]

    for q_num, q_text, ans, choices in q7_items:
        qp = doc.add_paragraph()
        set_para_format(qp, space_before=3, space_after=1, indent_left=0.3,
                        right_tab_at=RIGHT_MARGIN_TAB)
        add_run(qp, f"{q_num} {q_text}")
        if not student:
            add_run(qp, "\t")
            add_run(qp, ans, bold=True, boxed=True)
        for ci, choice in enumerate(choices, 1):
            cp = doc.add_paragraph()
            set_para_format(cp, space_before=0, space_after=0, indent_left=0.8)
            add_run(cp, f"　{ci}.　{choice}")
        doc.add_paragraph()

    # =========================================================
    # 大問８：Neo現代 Unit2「Britain」本文の穴埋め問題（10点）
    # =========================================================
    # 改ページ
    _pb_para = doc.add_paragraph()
    _pb_run = _pb_para.add_run()
    _br = OxmlElement('w:br')
    _br.set(qn('w:type'), 'page')
    _pb_run._r.append(_br)

    add_section_header(doc, 8,
        "次の英文を読み，（　１　）〜（　１０　）に入る最も適切な語を下の語群から選び，"
        "記号で答えなさい。ただし，同じ語は一度しか使えない。",
        points="１０")

    # 語群
    word_box_p = doc.add_paragraph()
    set_para_format(word_box_p, space_before=2, space_after=4, indent_left=0.3)
    add_run(word_box_p, "【語群】　", bold=True)
    add_run(word_box_p,
        "ア colonies　　イ culture　　ウ preserve　　エ famous　　オ factories\n"
        "カ ancient　　キ curry　　ク traditions　　ケ herbs　　コ Industrial\n"
        "※ サ image　　シ export　（使わないものが２つある）")

    # 本文（穴埋め）
    neo_paras = [
        ("　Britain has an （　１　） problem when it comes to food. Every other country seems to "
         "have its own special dishes: Italy has pizza and pasta, India has （　２　）, and France "
         "has haute cuisine. Britain, it seems, has nothing to offer."),
        ("　However, this is not entirely true. Britain was once （　３　） for its food. Back in "
         "the time of Queen Elizabeth I (1558–1603), people really knew how to eat well. Country "
         "houses had kitchens full of rosemary, parsley, thyme, and other （　４　）. Chefs used "
         "to travel around Europe to find new ideas and ingredients."),
        ("　Things changed dramatically around the end of the 19th century. The （　５　） "
         "Revolution forced many families to move from the countryside into crowded new cities "
         "like Manchester and Birmingham. Men, women, and even children worked long hours in "
         "（　６　）. They no longer had the time or energy to cook properly at home."),
        ("　In the 20th century, many people came to the UK from Britain's former （　７　） — "
         "including India, Pakistan, West Africa, the Caribbean, and Hong Kong. Other immigrants, "
         "such as the Irish, Jews, and Italians, were already living in Britain. These people did "
         "not simply give up their own ways of eating. They tried to （　８　） their food "
         "（　９　）."),
        ("　Food is one of the most powerful parts of a （　１０　）. You may forget your language "
         "or stop wearing traditional clothes. But when you see and smell a dish that your "
         "grandmother used to make, you immediately feel connected to your roots."),
    ]

    for text in neo_paras:
        p = doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        space_before=0, space_after=2, line_spacing=18)
        add_run(p, text)

    # 解答
    if not student:
        neo_ans = doc.add_paragraph()
        set_para_format(neo_ans, space_before=4, space_after=4, indent_left=0.3)
        add_run(neo_ans,
            "１サ　２キ　３エ　４ケ　５コ　６オ　７ア　８ウ　９ク　１０イ",
            bold=True, yellow=True)

    # =========================================================
    # 保存
    # =========================================================
    import os
    os.makedirs("/home/user/teacher-automation-lab/output", exist_ok=True)
    if student:
        out_path = "/home/user/teacher-automation-lab/output/exam_comm2_student.docx"
    else:
        out_path = "/home/user/teacher-automation-lab/output/exam_comm2_final.docx"
    doc.save(out_path)
    print(f"保存完了: {out_path}")

if __name__ == "__main__":
    build_exam(student=False)   # 教員用（解答入り）
    build_exam(student=True)    # 生徒配布用（解答なし）
