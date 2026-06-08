from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PRICES = {
    'タフブレーカー上':       8663,
    'タフブレーカー下':       6875,
    'ラグビージャージ黒':     6848,
    'ラグビージャージ赤':     6848,
    'スウェットパーカー':     7508,
    'スウェットパンツ':       5803,
    'ポロシャツ':             5033,
    'ハーフパンツ':           3212,
    'チームバッグ':           9845,
    'ショルダーガード':       6188,
    'ヘッドギア':             5198,
    'チームソックス（えんじ）': 1870,
    'チームソックス（紫）':   1870,
    'ラグビーパンツ（セプター）': 4510,
}

# ショルダーガードサイズ（フォーム回答結果 2026/05/26-30）
SG_SIZES = {
    29: 'S',   # 竹原浩輝
    36: 'M',   # 野村奏太
    38: 'L',   # 藤野裕心
    39: 'S',   # 須藤海翔
    40: 'M',   # 中所龍之介
    41: 'XL',  # 吉田青空
    42: 'S',   # 阿部浩直
    44: 'M',   # 田口輝真
    45: 'L',   # 遠藤将来（未回答・据え置き）
    47: 'L',   # 増喜晴也
}

def sock_size(s):
    s = str(s).replace('cm','').strip()
    return s + 'cm'

STUDENTS = [
    dict(num=29, name='竹原 浩輝', klass='A', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','2XL',1),('ラグビージャージ赤','2XL',1),
        ('スウェットパーカー','2XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','2XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[29],1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27),1),
        ('チームソックス（紫）',sock_size(27),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=30, name='葛西 律紀', klass='I', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','2XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','2XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=31, name='大津 究士', klass='D', items=[
        ('タフブレーカー上','2XL',1),('タフブレーカー下','2XL',1),
        ('ラグビージャージ黒','2XL',1),('ラグビージャージ赤','2XL',1),
        ('スウェットパーカー','2XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','3XL',1),('ハーフパンツ','2XL',1),
        ('チームバッグ','—',1),
        ('ヘッドギア','O',1),
        ('チームソックス（えんじ）',sock_size(29),1),
        ('チームソックス（紫）',sock_size(29),1),
        ('ラグビーパンツ（セプター）','XL',1)]),
    dict(num=32, name='稲垣 飛日人', klass='A', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','L',1),
        ('チームバッグ','—',1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(26.5),1),
        ('チームソックス（紫）',sock_size(26.5),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=33, name='岸田 涼', klass='D', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','L',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(26),1),
        ('チームソックス（紫）',sock_size(26),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=34, name='植作 渉', klass='A', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27),1),
        ('チームソックス（紫）',sock_size(27),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=35, name='粕谷 雄大', klass='F', items=[
        ('タフブレーカー上','2XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','2XL',1),('ラグビージャージ赤','2XL',1),
        ('スウェットパーカー','2XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','2XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','XL',2)]),  # 2枚
    dict(num=36, name='野村 奏太', klass='D', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[36],1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=37, name='齋藤 賢青', klass='H', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','L',1),
        ('チームバッグ','—',1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=38, name='藤野 裕心', klass='D', items=[
        ('タフブレーカー上','2XL',1),('タフブレーカー下','2XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[38],1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','XL',1)]),
    dict(num=39, name='須藤 海翔', klass='E', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','2XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','2XL',1),('スウェットパンツ','2XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[39],1),
        ('ヘッドギア','O',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','XL',1)]),
    dict(num=40, name='中所 龍之介', klass='K', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[40],1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=41, name='吉田 青空', klass='E', items=[
        ('タフブレーカー上','3XL',1),('タフブレーカー下','3XL',1),
        ('ラグビージャージ黒','3XL',1),('ラグビージャージ赤','3XL',1),
        ('スウェットパーカー','3XL',1),('スウェットパンツ','3XL',1),
        ('ポロシャツ','3XL',1),('ハーフパンツ','3XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[41],1),
        ('ヘッドギア','O',1),
        ('チームソックス（えんじ）',sock_size(29),1),
        ('チームソックス（紫）',sock_size(29),1),
        ('ラグビーパンツ（セプター）','3XL',1)]),
    dict(num=42, name='阿部 浩直', klass='E', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','2XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','2XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[42],1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27.5),1),
        ('チームソックス（紫）',sock_size(27.5),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=44, name='田口 輝真', klass='J', items=[
        ('タフブレーカー上','2XL',1),('タフブレーカー下','2XL',1),
        ('ラグビージャージ黒','2XL',1),('ラグビージャージ赤','2XL',1),
        ('スウェットパーカー','2XL',1),('スウェットパンツ','2XL',1),
        ('ポロシャツ','2XL',1),('ハーフパンツ','XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[44],1),
        ('ヘッドギア','O',1),
        ('チームソックス（えんじ）',sock_size(28),1),
        ('チームソックス（紫）',sock_size(28),1),
        ('ラグビーパンツ（セプター）','XL',1)]),
    dict(num=45, name='遠藤 将来', klass='F', items=[
        ('タフブレーカー上','XL',1),('タフブレーカー下','XL',1),
        ('ラグビージャージ黒','XL',1),('ラグビージャージ赤','XL',1),
        ('スウェットパーカー','XL',1),('スウェットパンツ','XL',1),
        ('ポロシャツ','XL',1),('ハーフパンツ','L',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[45],1),
        ('ヘッドギア','L',1),
        ('チームソックス（えんじ）',sock_size(27),1),
        ('チームソックス（紫）',sock_size(27),1),
        ('ラグビーパンツ（セプター）','L',1)]),
    dict(num=47, name='増喜 晴也', klass='E', items=[
        ('タフブレーカー上','2XL',1),('タフブレーカー下','2XL',1),
        ('ラグビージャージ黒','2XL',1),('ラグビージャージ赤','2XL',1),
        ('スウェットパーカー','2XL',1),('スウェットパンツ','2XL',1),
        ('ポロシャツ','2XL',1),('ハーフパンツ','2XL',1),
        ('チームバッグ','—',1),('ショルダーガード',SG_SIZES[47],1),
        ('ヘッドギア','O',1),
        ('チームソックス（えんじ）',sock_size(27),1),
        ('チームソックス（紫）',sock_size(27),1),
        ('ラグビーパンツ（セプター）','2XL',1)]),
]

def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '游明朝'
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_run(para, text, size=10.5, bold=False, color=None):
    r = para.add_run(text)
    set_font(r, size, bold, color)
    return r

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None, color=None):
    if bg:
        set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    add_run(p, text, size=size, bold=bold, color=color)

def add_student_page(doc, s, first=False):
    if not first:
        doc.add_page_break()

    body = doc.add_section() if False else None  # no-op

    # ── ヘッダー ───────────────────────────────
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(4)
    add_run(h, '令和８年度　駒澤大学高等学校ラグビー部', size=11, bold=True,
            color=(0x1f, 0x49, 0x7d))
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after  = Pt(6)
    add_run(h2, '新入部員　用具購入のお知らせ', size=14, bold=True,
            color=(0x1f, 0x49, 0x7d))

    # ── 宛名 ──────────────────────────────────
    addr = doc.add_paragraph()
    addr.paragraph_format.space_before = Pt(0)
    addr.paragraph_format.space_after  = Pt(2)
    add_run(addr, f'１年{s["klass"]}組　{s["num"]}番　{s["name"]}　保護者　様', size=11)

    # ── 訂正・お詫び notice ────────────────────
    notice_tbl = doc.add_table(rows=1, cols=1)
    notice_tbl.style = 'Table Grid'
    notice_cell = notice_tbl.cell(0, 0)
    set_cell_bg(notice_cell, 'FFF2CC')
    notice_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    np0 = notice_cell.paragraphs[0]
    np0.paragraph_format.space_before = Pt(3)
    np0.paragraph_format.space_after  = Pt(1)
    add_run(np0, '【訂正・お詫び】', size=10.5, bold=True)

    notice_lines = [
        '先日お渡しした案内書について、下記の通り訂正いたします。ご迷惑をおかけし、誠に申し訳ございませんでした。',
        '① ショルダーガードのサイズを修正しました。今一度ご確認ください。',
        '② お支払いについて：三菱UFJ銀行へのお振込みは不要です。代金はミズノ担当者より各ご家庭へ個別にご連絡いたします。',
    ]
    for line in notice_lines:
        np = notice_cell.add_paragraph()
        np.paragraph_format.space_before = Pt(1)
        np.paragraph_format.space_after  = Pt(1)
        if line.startswith('①') or line.startswith('②'):
            np.paragraph_format.left_indent = Pt(10)
        add_run(np, line, size=10)
    notice_cell.add_paragraph().paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 挨拶文 ────────────────────────────────
    greet = doc.add_paragraph()
    greet.paragraph_format.first_line_indent = Pt(10.5)
    greet.paragraph_format.space_before = Pt(0)
    greet.paragraph_format.space_after  = Pt(6)
    add_run(greet,
        '平素よりラグビー部の活動にご理解・ご協力をいただき、誠にありがとうございます。'
        '下記の通り、用具一式のご注文内容と金額をお知らせいたします。'
        '内容をご確認のうえ、ミズノ担当者からのご連絡をお待ちください。',
        size=10.5)

    # ── 品目テーブル ──────────────────────────
    col_w = [Cm(6.5), Cm(2.0), Cm(1.3), Cm(2.5), Cm(2.7)]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, cell in enumerate(tbl.rows[0].cells):
        cell.width = col_w[i]
    headers = ['品　名', 'サイズ', '数量', '単価（税込）', '小計（税込）']
    for i, (cell, h) in enumerate(zip(tbl.rows[0].cells, headers)):
        cell_text(cell, h, size=10, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, bg='1F497D',
                  color=(0xFF, 0xFF, 0xFF))

    total_calc = 0
    for idx, (name, size, qty) in enumerate(s['items']):
        unit = PRICES[name]
        sub  = unit * qty
        total_calc += sub
        row = tbl.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = col_w[i]
        bg = 'DEEAF1' if idx % 2 == 0 else None
        cell_text(row.cells[0], name,  size=10, bg=bg)
        cell_text(row.cells[1], size,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        cell_text(row.cells[2], str(qty), size=10, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        cell_text(row.cells[3], f'¥{unit:,}', size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, bg=bg)
        cell_text(row.cells[4], f'¥{sub:,}',  size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, bg=bg)

    # 合計行
    total_row = tbl.add_row()
    for i, cell in enumerate(total_row.cells):
        cell.width = col_w[i]
    cell_text(total_row.cells[0], '合　計', size=11, bold=True,
              align=WD_ALIGN_PARAGRAPH.RIGHT, bg='1F497D', color=(0xFF,0xFF,0xFF))
    total_row.cells[1].merge(total_row.cells[2]).merge(total_row.cells[3])
    set_cell_bg(total_row.cells[1], '1F497D')
    cell_text(total_row.cells[4], f'¥{total_calc:,}', size=11, bold=True,
              align=WD_ALIGN_PARAGRAPH.RIGHT, bg='1F497D', color=(0xFF,0xFF,0xFF))

    # ── お支払い案内 ──────────────────────────
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(8)
    sp.paragraph_format.space_after  = Pt(2)
    add_run(sp, '【お支払いについて】', size=10.5, bold=True)

    pay_lines = [
        '代金のお支払いは、ミズノ担当者より各ご家庭へ個別にご連絡いたします。',
        '三菱UFJ銀行へのお振込みは不要です。',
        '担当者からの連絡をお待ちください。',
    ]
    for line in pay_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Pt(14)
        add_run(p, line, size=10)

    # ── 署名 ──────────────────────────────────
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.paragraph_format.space_before = Pt(8)
    sig.paragraph_format.space_after  = Pt(0)
    add_run(sig, '駒澤大学高等学校　ラグビー部顧問　占部 涼也', size=10)

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height    = Cm(29.7)
    sec.page_width     = Cm(21.0)
    sec.top_margin     = Cm(2.0)
    sec.bottom_margin  = Cm(2.0)
    sec.left_margin    = Cm(2.5)
    sec.right_margin   = Cm(2.5)

    doc.styles['Normal'].font.name = '游明朝'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

    for i, s in enumerate(STUDENTS):
        if i > 0:
            doc.add_page_break()
        add_student_page(doc, s, first=(i == 0))

    os.makedirs('/home/user/teacher-automation-lab/output', exist_ok=True)
    out = '/home/user/teacher-automation-lab/output/令和8年度_新入生用具購入案内書_全員分.docx'
    doc.save(out)
    print(f'完了: {out}')

    # 合計金額確認
    for s in STUDENTS:
        calc = sum(PRICES[n] * q for n, sz, q in s['items'])
        match = '✅' if calc == s.get('total', calc) else f'❌ 期待値{s.get("total")}'
        print(f'{s["num"]} {s["name"]}: ¥{calc:,} {match}')

main()
